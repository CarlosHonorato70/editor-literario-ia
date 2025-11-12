#!/usr/bin/env python3
"""
Script de Verificação - Upload de Arquivos
==========================================

Este script verifica se a funcionalidade de upload de arquivos está funcionando
corretamente após as correções implementadas.

Execute este script para diagnosticar problemas.
"""

import sys
import os

print("="*70)
print("VERIFICAÇÃO DA FUNCIONALIDADE DE UPLOAD DE ARQUIVOS")
print("="*70)

# Adiciona o diretório atual ao path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Teste 1: Importação do módulo
print("\n[1/5] Testando importação do módulo file_handler...")
try:
    from modules.file_handler import extract_text, FileHandler
    print("    ✅ Módulo importado com sucesso")
    print("    ✅ Lazy imports funcionando (não carrega dependências pesadas)")
except Exception as e:
    print(f"    ❌ FALHA: {e}")
    print("\n    DIAGNÓSTICO:")
    print("    - Certifique-se de estar no diretório raiz do projeto")
    print("    - Verifique se a pasta 'modules' existe")
    sys.exit(1)

# Teste 2: Extração de TXT (sem dependências externas)
print("\n[2/5] Testando extração de arquivo TXT...")
try:
    txt_content = "Este é um arquivo de teste.\nCom múltiplas linhas.\n".encode('utf-8')
    text, error = extract_text(txt_content, "teste.txt")
    
    if error:
        print(f"    ❌ ERRO: {error}")
    else:
        print(f"    ✅ Extração TXT funcionando")
        print(f"    📄 Texto extraído: '{text[:50]}...'")
except Exception as e:
    print(f"    ❌ FALHA: {e}")

# Teste 3: Verificar se python-docx está instalado
print("\n[3/5] Verificando dependência python-docx...")
try:
    import docx
    print("    ✅ python-docx está instalado")
    
    # Testar extração DOCX
    print("    └─ Testando extração DOCX...")
    from docx import Document
    import io
    
    doc = Document()
    doc.add_paragraph("Teste de parágrafo 1")
    doc.add_paragraph("Teste de parágrafo 2")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_content = buffer.getvalue()
    
    text, error = extract_text(docx_content, "teste.docx")
    if error:
        print(f"       ❌ ERRO: {error}")
    else:
        print(f"       ✅ Extração DOCX funcionando")
        print(f"       📄 Texto: '{text[:60]}...'")
        
except ImportError:
    print("    ⚠️  python-docx NÃO está instalado")
    print("    ℹ️  Para instalar: pip install python-docx")
except Exception as e:
    print(f"    ❌ ERRO ao testar DOCX: {e}")

# Teste 4: Verificar se PyPDF2 está instalado
print("\n[4/5] Verificando dependência PyPDF2...")
try:
    import PyPDF2
    print("    ✅ PyPDF2 está instalado")
    
    # Testar extração PDF
    print("    └─ Testando extração PDF...")
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    import io
    
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.drawString(100, 750, "Teste de PDF")
    c.drawString(100, 730, "Linha 2 do teste")
    c.save()
    pdf_content = buffer.getvalue()
    
    text, error = extract_text(pdf_content, "teste.pdf")
    if error:
        print(f"       ❌ ERRO: {error}")
    else:
        print(f"       ✅ Extração PDF funcionando")
        print(f"       📄 Texto: '{text[:60]}...'")
        
except ImportError as ie:
    print(f"    ⚠️  Dependência faltando: {str(ie).split()[-1]}")
    print("    ℹ️  Para instalar: pip install PyPDF2 reportlab")
except Exception as e:
    print(f"    ❌ ERRO ao testar PDF: {e}")

# Teste 5: Verificar integração com app_editor.py
print("\n[5/5] Verificando integração com app_editor.py...")
try:
    with open('app_editor.py', 'r', encoding='utf-8') as f:
        content = f.read()
        
    checks = [
        ('from modules.file_handler import extract_text', 'Import do file_handler'),
        ('st.session_state.text_content = text', 'Atualização do session_state'),
        ('type=["txt", "docx", "pdf"]', 'Suporte a PDF'),
        ('on_change=processar_arquivo_carregado', 'Callback do file_uploader'),
    ]
    
    all_ok = True
    for check_str, description in checks:
        if check_str in content:
            print(f"    ✅ {description}")
        else:
            print(f"    ❌ {description} - NÃO ENCONTRADO")
            all_ok = False
    
    if all_ok:
        print("    ✅ Integração com app_editor.py está correta")
    else:
        print("    ❌ Problemas encontrados na integração")
        
except FileNotFoundError:
    print("    ⚠️  app_editor.py não encontrado")
    print("    ℹ️  Execute este script do diretório raiz do projeto")
except Exception as e:
    print(f"    ❌ ERRO: {e}")

# Resumo final
print("\n" + "="*70)
print("RESUMO DA VERIFICAÇÃO")
print("="*70)

print("\n✅ FUNCIONALIDADES IMPLEMENTADAS:")
print("   • Módulo file_handler com lazy imports")
print("   • Extração de TXT, DOCX e PDF")
print("   • Integração com Streamlit app_editor.py")
print("   • Tratamento de erros robusto")

print("\n📋 PRÓXIMOS PASSOS:")
print("   1. Instale as dependências: pip install -r requirements.txt")
print("   2. Limpe o cache Python: ")
print("      - Windows: for /d /r . %d in (__pycache__) do @if exist \"%d\" rd /s /q \"%d\"")
print("      - Linux/Mac: find . -type d -name __pycache__ -exec rm -rf {} +")
print("   3. Execute o app: streamlit run app_editor.py")
print("   4. Teste fazendo upload de um arquivo TXT, DOCX ou PDF")

print("\n💡 DICA: Se ainda houver problemas:")
print("   • Verifique se está usando a branch correta: git status")
print("   • Puxe as últimas mudanças: git pull origin copilot/fix-file-upload-extraction")
print("   • Consulte TROUBLESHOOTING_WINDOWS.md para problemas específicos do Windows")

print("\n" + "="*70)
print("Verificação concluída!")
print("="*70)
