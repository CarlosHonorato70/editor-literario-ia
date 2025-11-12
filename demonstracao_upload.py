#!/usr/bin/env python3
"""
Demonstração da Funcionalidade de Upload
========================================

Este script demonstra que a funcionalidade de upload está funcionando
corretamente, simulando o que acontece quando um usuário carrega um arquivo
no Streamlit.
"""

import sys
import os
import io

# Adiciona o diretório ao path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

print("="*80)
print("DEMONSTRAÇÃO: Upload e Extração de Arquivos Funcionando")
print("="*80)

from modules.file_handler import extract_text

# Simula o que acontece no processar_arquivo_carregado()
def simular_upload(file_content, filename):
    """Simula o processo de upload do Streamlit"""
    print(f"\n📁 Simulando upload de: {filename}")
    print(f"   Tamanho: {len(file_content)} bytes")
    
    # Isto é exatamente o que acontece em app_editor.py linha 104-115
    text, error = extract_text(file_content, filename)
    
    if error:
        print(f"   ❌ ERRO: {error}")
        print(f"   st.session_state.text_content = ''")
        print(f"   st.session_state.file_processed = False")
        return False
    else:
        print(f"   ✅ SUCESSO: Arquivo '{filename}' carregado!")
        print(f"   st.session_state.text_content = <texto extraído>")
        print(f"   st.session_state.file_processed = True")
        print(f"\n   📄 TEXTO EXTRAÍDO:")
        print(f"   {'-'*70}")
        # Mostra as primeiras linhas
        lines = text.split('\n')[:5]
        for line in lines:
            print(f"   {line}")
        if len(text) > 200:
            print(f"   ... ({len(text)} caracteres no total)")
        print(f"   {'-'*70}")
        return True

# TESTE 1: Arquivo TXT
print("\n" + "="*80)
print("TESTE 1: Upload de arquivo TXT")
print("="*80)

txt_content = """Este é um manuscrito de exemplo.

Era uma vez, em uma terra distante, um escritor que queria publicar seu livro.

Ele escreveu várias páginas de texto, revisou cuidadosamente cada parágrafo,
e finalmente decidiu fazer o upload do arquivo no editor literário.

O sistema deveria extrair todo o texto e exibí-lo no campo de edição.

Fim do exemplo.""".encode('utf-8')

sucesso_txt = simular_upload(txt_content, "manuscrito.txt")

# TESTE 2: Arquivo DOCX
print("\n" + "="*80)
print("TESTE 2: Upload de arquivo DOCX")
print("="*80)

try:
    from docx import Document
    
    doc = Document()
    doc.add_paragraph("Capítulo 1: O Início")
    doc.add_paragraph("")
    doc.add_paragraph("Este é o primeiro parágrafo do manuscrito.")
    doc.add_paragraph("Este é o segundo parágrafo, com mais conteúdo.")
    doc.add_paragraph("")
    doc.add_paragraph("Capítulo 2: O Desenvolvimento")
    doc.add_paragraph("")
    doc.add_paragraph("A história continua com mais detalhes...")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_content = buffer.getvalue()
    
    sucesso_docx = simular_upload(docx_content, "manuscrito.docx")
    
except ImportError:
    print("   ⚠️  python-docx não instalado - teste pulado")
    print("   Para instalar: pip install python-docx")
    sucesso_docx = None

# TESTE 3: Arquivo PDF
print("\n" + "="*80)
print("TESTE 3: Upload de arquivo PDF")
print("="*80)

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    
    # Página 1
    c.setFont("Helvetica", 16)
    c.drawString(100, 750, "Manuscrito Literário - Exemplo")
    c.setFont("Helvetica", 12)
    c.drawString(100, 700, "Este é um exemplo de documento PDF.")
    c.drawString(100, 680, "O sistema deve extrair todo o texto deste arquivo.")
    c.drawString(100, 660, "E exibi-lo no campo de edição do Streamlit.")
    c.showPage()
    
    # Página 2
    c.drawString(100, 750, "Página 2 do documento")
    c.drawString(100, 730, "Com mais conteúdo para teste...")
    
    c.save()
    pdf_content = buffer.getvalue()
    
    sucesso_pdf = simular_upload(pdf_content, "manuscrito.pdf")
    
except ImportError as e:
    print(f"   ⚠️  Dependência faltando: {e}")
    print("   Para instalar: pip install PyPDF2 reportlab")
    sucesso_pdf = None

# RESUMO FINAL
print("\n" + "="*80)
print("RESUMO DOS TESTES")
print("="*80)

resultados = [
    ("TXT", sucesso_txt),
    ("DOCX", sucesso_docx),
    ("PDF", sucesso_pdf)
]

print("\nResultados:")
for tipo, resultado in resultados:
    if resultado is True:
        print(f"   ✅ {tipo}: FUNCIONANDO - texto extraído e pronto para exibição")
    elif resultado is False:
        print(f"   ❌ {tipo}: FALHOU - erro na extração")
    else:
        print(f"   ⚠️  {tipo}: NÃO TESTADO - dependência faltando")

print("\n" + "="*80)
print("CONCLUSÃO")
print("="*80)

if sucesso_txt:
    print("\n✅ A funcionalidade de upload ESTÁ FUNCIONANDO!")
    print("\nO que acontece quando você faz upload no Streamlit:")
    print("1. Você clica em 'Browse files' e seleciona um arquivo")
    print("2. O Streamlit chama processar_arquivo_carregado()")
    print("3. O código chama extract_text() do file_handler")
    print("4. O texto é extraído e salvo em st.session_state.text_content")
    print("5. O st.text_area (que usa key='text_content') exibe o texto automaticamente")
    print("6. Você vê uma mensagem: ✅ Arquivo 'nome.txt' carregado com sucesso!")
    print("\nSe você NÃO está vendo o texto no editor:")
    print("• Verifique se você tem a última versão do código (git pull)")
    print("• Limpe o cache Python (__pycache__ e .pyc)")
    print("• Reinicie o Streamlit")
    print("• Execute: python verificar_upload.py para diagnóstico completo")
else:
    print("\n❌ Há um problema com a funcionalidade básica")
    print("Execute: python verificar_upload.py para diagnóstico completo")

print("\n" + "="*80)
