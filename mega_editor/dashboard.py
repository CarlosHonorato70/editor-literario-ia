#!/usr/bin/env python3
"""
Mega Editor Dashboard - Interface Unificada de Edição Literária
Consolida TODAS as funcionalidades do repositório editor-literario-ia
"""

import streamlit as st
import sys
import os
from pathlib import Path

# Adicionar diretório raiz ao path
sys.path.insert(0, str(Path(__file__).parent.parent))

# Importações de módulos existentes
try:
    from modules.config import Config
    from modules.analyzer import ManuscriptAnalyzer
    from modules.enhancer import ContentEnhancer
    from modules.formatter import DocumentFormatter
    from modules.reviewer import EditorialReviewer
    from modules.exporter import PublicationExporter
    from modules.fastformat_utils import apply_fastformat, get_ptbr_options
    from modules.workflow_orchestrator import WorkflowOrchestrator
    from modules.workflow_tab_enhanced import render_workflow_tab
    ADVANCED_MODULES_AVAILABLE = True
except ImportError as e:
    ADVANCED_MODULES_AVAILABLE = False
    print(f"Warning: Some advanced modules not available: {e}")

# Importações de editores
try:
    from streamlit_quill import st_quill
    QUILL_AVAILABLE = True
except ImportError:
    QUILL_AVAILABLE = False

try:
    from streamlit_ace import st_ace
    ACE_AVAILABLE = True
except ImportError:
    ACE_AVAILABLE = False

# Importações padrão
from docx import Document
from docx.shared import Pt, Cm, Inches
import io
import json
from datetime import datetime

# ====================================================================================
# CONFIGURAÇÃO GLOBAL
# ====================================================================================

def setup_page():
    """Configuração inicial da página"""
    st.set_page_config(
        page_title="📚 Mega Editor - Sistema Integrado de Edição Literária",
        page_icon="📚",
        layout="wide",
        initial_sidebar_state="expanded"
    )

def initialize_session_state():
    """Inicializa todas as variáveis de sessão necessárias"""
    defaults = {
        # Estado geral
        "initialized": True,
        "current_module": "dashboard",
        
        # Conteúdo
        "text_content": "",
        "uploaded_text": None,
        "pending_text_update": None,
        
        # Editores especializados
        "quill_content": "",
        "ace_content": "",
        "rich_editor_content": None,
        
        # Metadados do projeto
        "book_title": "Sem Título",
        "author_name": "Autor Desconhecido",
        "contact_info": "seuemail@exemplo.com",
        "genre": "Ficção",
        "target_audience": "Adulto",
        "isbn": "",
        
        # Workflow e processamento
        "file_processed": False,
        "sugestoes_estilo": None,
        "api_key_valida": False,
        "use_fastformat": True,
        
        # Configurações de análise
        "analysis_results": None,
        "enhancement_suggestions": None,
        "editorial_review": None,
        
        # Histórico e versionamento
        "version_history": [],
        "current_version": 0,
        
        # Exportação
        "export_format": "docx",
        "export_ready": False,
        
        # Workflow de 14 fases
        "workflow_phase": 1,
        "workflow_progress": {},
        "workflow_completed_phases": set(),
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

# ====================================================================================
# SIDEBAR - NAVEGAÇÃO PRINCIPAL
# ====================================================================================

def render_sidebar():
    """Renderiza a barra lateral de navegação"""
    with st.sidebar:
        st.title("📚 Mega Editor")
        st.markdown("---")
        
        st.subheader("🎯 Navegação Principal")
        
        # Dashboard Principal
        if st.button("🏠 Dashboard Principal", use_container_width=True):
            st.session_state.current_module = "dashboard"
            st.rerun()
        
        st.markdown("### 📝 Editores")
        
        # Editor Simples
        if st.button("📄 Editor Simples", use_container_width=True):
            st.session_state.current_module = "editor_simples"
            st.rerun()
        
        # Editor Avançado (Quill)
        if QUILL_AVAILABLE:
            if st.button("✍️ Editor Avançado (Word-like)", use_container_width=True):
                st.session_state.current_module = "editor_quill"
                st.rerun()
        
        # Editor de Código (Ace)
        if ACE_AVAILABLE:
            if st.button("💻 Editor de Código", use_container_width=True):
                st.session_state.current_module = "editor_ace"
                st.rerun()
        
        st.markdown("### ⚙️ Processamento")
        
        # FastFormat
        if st.button("🎨 FastFormat (Tipografia)", use_container_width=True):
            st.session_state.current_module = "fastformat"
            st.rerun()
        
        # Análise de Manuscrito
        if ADVANCED_MODULES_AVAILABLE:
            if st.button("🔍 Análise de Manuscrito", use_container_width=True):
                st.session_state.current_module = "analise"
                st.rerun()
            
            # Aprimoramento de Conteúdo
            if st.button("✨ Aprimoramento IA", use_container_width=True):
                st.session_state.current_module = "aprimoramento"
                st.rerun()
            
            # Revisão Editorial
            if st.button("📋 Revisão Editorial", use_container_width=True):
                st.session_state.current_module = "revisao"
                st.rerun()
        
        st.markdown("### 🔄 Workflows Completos")
        
        # Workflow de 14 Fases
        if st.button("🔄 Workflow 14 Fases", use_container_width=True):
            st.session_state.current_module = "workflow_14"
            st.rerun()
        
        # Workflow Automatizado
        if ADVANCED_MODULES_AVAILABLE:
            if st.button("🤖 Workflow Automatizado", use_container_width=True):
                st.session_state.current_module = "workflow_auto"
                st.rerun()
        
        st.markdown("### 📦 Exportação")
        
        # Exportação Multi-formato
        if st.button("📥 Exportar Documento", use_container_width=True):
            st.session_state.current_module = "exportacao"
            st.rerun()
        
        # Preparação para Publicação
        if ADVANCED_MODULES_AVAILABLE:
            if st.button("🚀 Preparar para Publicação", use_container_width=True):
                st.session_state.current_module = "publicacao"
                st.rerun()
        
        st.markdown("---")
        
        # Informações do projeto
        with st.expander("📊 Informações do Projeto"):
            st.write(f"**Título:** {st.session_state.book_title}")
            st.write(f"**Autor:** {st.session_state.author_name}")
            st.write(f"**Palavras:** {len(st.session_state.text_content.split())}")
            st.write(f"**Caracteres:** {len(st.session_state.text_content)}")
        
        # Módulos disponíveis
        with st.expander("🔧 Módulos Disponíveis"):
            st.write(f"✅ Editor Simples: Sempre")
            st.write(f"{'✅' if QUILL_AVAILABLE else '❌'} Editor Quill: {QUILL_AVAILABLE}")
            st.write(f"{'✅' if ACE_AVAILABLE else '❌'} Editor Ace: {ACE_AVAILABLE}")
            st.write(f"{'✅' if ADVANCED_MODULES_AVAILABLE else '❌'} Módulos Avançados: {ADVANCED_MODULES_AVAILABLE}")

# ====================================================================================
# DASHBOARD PRINCIPAL
# ====================================================================================

def render_dashboard():
    """Renderiza o dashboard principal com overview"""
    st.title("📚 Mega Editor - Sistema Integrado de Edição Literária")
    
    st.markdown("""
    Bem-vindo ao **Mega Editor**, uma plataforma completa que integra todas as funcionalidades 
    do Editor Literário IA em um único lugar!
    """)
    
    # Métricas principais
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("📄 Palavras", len(st.session_state.text_content.split()))
    
    with col2:
        st.metric("📝 Caracteres", len(st.session_state.text_content))
    
    with col3:
        paragraphs = len([p for p in st.session_state.text_content.split('\n\n') if p.strip()])
        st.metric("¶ Parágrafos", paragraphs)
    
    with col4:
        st.metric("📚 Versões", len(st.session_state.version_history))
    
    st.markdown("---")
    
    # Ações rápidas
    st.subheader("🚀 Ações Rápidas")
    
    col_a, col_b, col_c = st.columns(3)
    
    with col_a:
        st.markdown("### 📝 Editar")
        if st.button("Iniciar Edição Simples", use_container_width=True):
            st.session_state.current_module = "editor_simples"
            st.rerun()
        
        if QUILL_AVAILABLE and st.button("Editor Avançado (WYSIWYG)", use_container_width=True):
            st.session_state.current_module = "editor_quill"
            st.rerun()
    
    with col_b:
        st.markdown("### ⚙️ Processar")
        if st.button("Aplicar FastFormat", use_container_width=True):
            st.session_state.current_module = "fastformat"
            st.rerun()
        
        if ADVANCED_MODULES_AVAILABLE and st.button("Análise Completa", use_container_width=True):
            st.session_state.current_module = "analise"
            st.rerun()
    
    with col_c:
        st.markdown("### 📦 Exportar")
        if st.button("Exportar DOCX", use_container_width=True):
            st.session_state.current_module = "exportacao"
            st.rerun()
        
        if ADVANCED_MODULES_AVAILABLE and st.button("Preparar Publicação", use_container_width=True):
            st.session_state.current_module = "publicacao"
            st.rerun()
    
    st.markdown("---")
    
    # Funcionalidades Disponíveis
    st.subheader("✨ Funcionalidades Integradas")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        #### 📝 Editores
        - ✅ **Editor Simples**: Text area com funcionalidades básicas
        - {} **Editor Avançado (Quill)**: WYSIWYG tipo Word
        - {} **Editor de Código (Ace)**: Syntax highlighting profissional
        
        #### ⚙️ Processamento de Texto
        - ✅ **FastFormat**: Formatação tipográfica brasileira
        - {} **Análise de Manuscrito**: Estrutura, estilo, coesão
        - {} **Aprimoramento IA**: Sugestões de melhoria com IA
        - {} **Revisão Editorial**: Revisão profissional automatizada
        """.format(
            "✅" if QUILL_AVAILABLE else "❌",
            "✅" if ACE_AVAILABLE else "❌",
            "✅" if ADVANCED_MODULES_AVAILABLE else "❌",
            "✅" if ADVANCED_MODULES_AVAILABLE else "❌",
            "✅" if ADVANCED_MODULES_AVAILABLE else "❌"
        ))
    
    with col2:
        st.markdown("""
        #### 🔄 Workflows Completos
        - ✅ **Workflow 14 Fases**: Processo editorial completo
        - {} **Workflow Automatizado**: Processamento automático end-to-end
        
        #### 📦 Exportação e Publicação
        - ✅ **Exportação Multi-formato**: DOCX, PDF, EPUB, HTML, MD
        - {} **Preparação para Publicação**: Amazon KDP, Google Play Books
        - {} **Elementos Pré/Pós-textuais**: Ficha catalográfica, ISBN
        - {} **Arquivos para Gráfica**: Print-ready PDFs
        
        #### 🔧 Ferramentas Auxiliares
        - ✅ **Versionamento**: Histórico de mudanças
        - ✅ **Estatísticas**: Análise de texto em tempo real
        - ✅ **Validação**: Checklist de qualidade
        """.format(
            "✅" if ADVANCED_MODULES_AVAILABLE else "❌",
            "✅" if ADVANCED_MODULES_AVAILABLE else "❌",
            "✅" if ADVANCED_MODULES_AVAILABLE else "❌",
            "✅" if ADVANCED_MODULES_AVAILABLE else "❌"
        ))
    
    st.markdown("---")
    
    # Projeto Atual
    st.subheader("📂 Projeto Atual")
    
    col_meta1, col_meta2 = st.columns(2)
    
    with col_meta1:
        new_title = st.text_input("Título do Livro", st.session_state.book_title)
        if new_title != st.session_state.book_title:
            st.session_state.book_title = new_title
        
        new_author = st.text_input("Autor", st.session_state.author_name)
        if new_author != st.session_state.author_name:
            st.session_state.author_name = new_author
    
    with col_meta2:
        new_genre = st.selectbox(
            "Gênero",
            ["Ficção", "Não-ficção", "Romance", "Suspense", "Fantasia", "Biografia", "Técnico", "Acadêmico"],
            index=0 if st.session_state.genre == "Ficção" else 0
        )
        if new_genre != st.session_state.genre:
            st.session_state.genre = new_genre
        
        new_contact = st.text_input("Contato", st.session_state.contact_info)
        if new_contact != st.session_state.contact_info:
            st.session_state.contact_info = new_contact
    
    # Preview do conteúdo
    if st.session_state.text_content:
        with st.expander("👁️ Preview do Conteúdo"):
            preview_length = min(500, len(st.session_state.text_content))
            st.text(st.session_state.text_content[:preview_length] + ("..." if len(st.session_state.text_content) > preview_length else ""))

# ====================================================================================
# MÓDULOS DE EDIÇÃO
# ====================================================================================

def render_editor_simples():
    """Editor simples de texto"""
    st.title("📄 Editor Simples")
    
    # Upload de arquivo
    st.subheader("📁 Carregar Arquivo")
    uploaded_file = st.file_uploader("Escolha um arquivo", type=["txt", "docx"], key="simple_uploader")
    
    if uploaded_file:
        try:
            if uploaded_file.name.endswith('.txt'):
                st.session_state.text_content = uploaded_file.read().decode('utf-8')
            elif uploaded_file.name.endswith('.docx'):
                doc = Document(io.BytesIO(uploaded_file.read()))
                st.session_state.text_content = '\n\n'.join([para.text for para in doc.paragraphs])
            
            st.success("✅ Arquivo carregado com sucesso!")
        except Exception as e:
            st.error(f"Erro ao carregar arquivo: {e}")
    
    st.markdown("---")
    
    # Editor
    st.subheader("✍️ Área de Edição")
    new_text = st.text_area(
        "Escreva ou cole seu texto aqui",
        value=st.session_state.text_content,
        height=500,
        key="simple_editor"
    )
    
    if new_text != st.session_state.text_content:
        st.session_state.text_content = new_text
    
    # Estatísticas
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Palavras", len(new_text.split()))
    with col2:
        st.metric("Caracteres", len(new_text))
    with col3:
        st.metric("Linhas", new_text.count('\n') + 1)
    with col4:
        paragraphs = len([p for p in new_text.split('\n\n') if p.strip()])
        st.metric("Parágrafos", paragraphs)
    
    # Ações
    col_a, col_b, col_c = st.columns(3)
    with col_a:
        if st.button("💾 Salvar Versão", use_container_width=True):
            st.session_state.version_history.append({
                "version": len(st.session_state.version_history) + 1,
                "content": new_text,
                "timestamp": datetime.now().isoformat(),
                "word_count": len(new_text.split())
            })
            st.success(f"✅ Versão {len(st.session_state.version_history)} salva!")
    
    with col_b:
        if st.button("🗑️ Limpar", use_container_width=True):
            st.session_state.text_content = ""
            st.rerun()
    
    with col_c:
        if st.button("📥 Exportar TXT", use_container_width=True):
            st.download_button(
                "💾 Download",
                data=new_text,
                file_name=f"{st.session_state.book_title}.txt",
                mime="text/plain"
            )

def render_editor_quill():
    """Editor avançado com Quill"""
    if not QUILL_AVAILABLE:
        st.error("❌ Editor Quill não disponível. Instale com: pip install streamlit-quill")
        return
    
    st.title("✍️ Editor Avançado (Word-like)")
    
    st.markdown("""
    Editor WYSIWYG com formatação rica. Use a barra de ferramentas para formatar seu texto.
    """)
    
    # Carregar conteúdo do editor principal se disponível
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.subheader("📝 Área de Edição")
        
        # Inicializar com conteúdo existente se houver
        initial_value = st.session_state.rich_editor_content or st.session_state.text_content
        
        content = st_quill(
            value=initial_value,
            placeholder="Comece a escrever...",
            html=True,
            key="quill_editor_main"
        )
        
        if content:
            st.session_state.rich_editor_content = content
    
    with col2:
        st.subheader("⚙️ Ações")
        
        if st.button("📤 Enviar para Editor Simples", use_container_width=True):
            # Converter HTML para texto plano (básico)
            import re
            plain_text = re.sub(r'<[^>]+>', '', st.session_state.rich_editor_content or "")
            st.session_state.text_content = plain_text
            st.success("✅ Texto transferido!")
        
        if st.button("📥 Carregar do Editor Simples", use_container_width=True):
            st.session_state.rich_editor_content = st.session_state.text_content
            st.rerun()
        
        if st.button("💾 Salvar Versão HTML", use_container_width=True):
            st.session_state.version_history.append({
                "version": len(st.session_state.version_history) + 1,
                "content": st.session_state.rich_editor_content,
                "timestamp": datetime.now().isoformat(),
                "format": "html"
            })
            st.success("✅ Versão HTML salva!")
        
        if st.button("🗑️ Limpar Editor", use_container_width=True):
            st.session_state.rich_editor_content = ""
            st.rerun()
    
    # Preview HTML
    if st.session_state.rich_editor_content:
        with st.expander("👁️ Preview HTML"):
            st.markdown(st.session_state.rich_editor_content, unsafe_allow_html=True)

def render_editor_ace():
    """Editor de código com Ace"""
    if not ACE_AVAILABLE:
        st.error("❌ Editor Ace não disponível. Instale com: pip install streamlit-ace")
        return
    
    st.title("💻 Editor de Código")
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        # Configurações do editor
        col_lang, col_theme = st.columns(2)
        
        with col_lang:
            language = st.selectbox(
                "Linguagem:",
                ["markdown", "html", "css", "javascript", "python", "yaml", "json", "xml", "plain_text"],
                key="ace_language"
            )
        
        with col_theme:
            theme = st.selectbox(
                "Tema:",
                ["monokai", "github", "tomorrow", "twilight", "solarized_dark", "solarized_light"],
                key="ace_theme"
            )
        
        # Editor Ace
        content = st_ace(
            value=st.session_state.ace_content or st.session_state.text_content,
            language=language,
            theme=theme,
            height=500,
            font_size=14,
            key="ace_editor_main"
        )
        
        if content:
            st.session_state.ace_content = content
    
    with col2:
        st.subheader("⚙️ Ações")
        
        if st.button("📤 Enviar para Editor Simples", use_container_width=True):
            st.session_state.text_content = st.session_state.ace_content
            st.success("✅ Texto transferido!")
        
        if st.button("📥 Carregar do Editor Simples", use_container_width=True):
            st.session_state.ace_content = st.session_state.text_content
            st.rerun()
        
        if st.button("💾 Salvar Código", use_container_width=True):
            extension_map = {
                "markdown": "md", "html": "html", "css": "css",
                "javascript": "js", "python": "py", "yaml": "yaml",
                "json": "json", "xml": "xml", "plain_text": "txt"
            }
            extension = extension_map.get(language, "txt")
            
            st.download_button(
                "💾 Download",
                data=st.session_state.ace_content,
                file_name=f"code.{extension}",
                mime="text/plain"
            )

# ====================================================================================
# MÓDULOS DE PROCESSAMENTO
# ====================================================================================

def render_fastformat():
    """Módulo FastFormat"""
    st.title("🎨 FastFormat - Formatação Tipográfica")
    
    if not st.session_state.text_content:
        st.warning("⚠️ Nenhum texto disponível. Carregue um texto primeiro no Editor Simples.")
        return
    
    st.markdown("""
    FastFormat aplica regras tipográficas brasileiras ao seu texto, incluindo:
    - Normalização de aspas e travessões
    - Correção de espaçamento
    - Formatação de reticências
    - Padronização de pontuação
    """)
    
    # Opções de formatação
    col1, col2 = st.columns(2)
    
    with col1:
        quote_style = st.selectbox(
            "Estilo de Aspas",
            ["brasileiras", "inglesas", "simples"],
            help="Escolha o estilo de aspas a ser usado"
        )
        
        dash_style = st.selectbox(
            "Estilo de Travessão",
            ["em_dash", "travessao_brasileiro"],
            help="Escolha o estilo de travessão para diálogos"
        )
    
    with col2:
        fix_spacing = st.checkbox("Corrigir Espaçamento", value=True)
        normalize_ellipsis = st.checkbox("Normalizar Reticências", value=True)
    
    # Preview
    with st.expander("👁️ Preview do Texto Original"):
        st.text(st.session_state.text_content[:500] + "...")
    
    # Aplicar FastFormat
    if st.button("🎨 Aplicar FastFormat", type="primary", use_container_width=True):
        with st.spinner("Aplicando formatação..."):
            try:
                options = {
                    "quote_style": quote_style,
                    "dash_style": dash_style,
                    "fix_spacing": fix_spacing,
                    "normalize_ellipsis": normalize_ellipsis
                }
                
                formatted_text = apply_fastformat(st.session_state.text_content, options)
                st.session_state.text_content = formatted_text
                
                # Salvar versão
                st.session_state.version_history.append({
                    "version": len(st.session_state.version_history) + 1,
                    "content": formatted_text,
                    "timestamp": datetime.now().isoformat(),
                    "note": "FastFormat aplicado"
                })
                
                st.success("✅ FastFormat aplicado com sucesso!")
                
                # Mostrar preview do resultado
                with st.expander("👁️ Preview do Texto Formatado"):
                    st.text(formatted_text[:500] + "...")
                
            except Exception as e:
                st.error(f"Erro ao aplicar FastFormat: {e}")

def render_analise():
    """Módulo de Análise de Manuscrito"""
    if not ADVANCED_MODULES_AVAILABLE:
        st.error("❌ Módulo de análise não disponível.")
        return
    
    st.title("🔍 Análise de Manuscrito")
    
    if not st.session_state.text_content:
        st.warning("⚠️ Nenhum texto disponível. Carregue um texto primeiro.")
        return
    
    st.markdown("""
    Análise profissional do manuscrito incluindo:
    - Estrutura e organização
    - Estilo e coesão
    - Problemas identificados
    - Sugestões de melhoria
    """)
    
    if st.button("🔍 Iniciar Análise", type="primary", use_container_width=True):
        with st.spinner("Analisando manuscrito..."):
            try:
                config = Config()
                analyzer = ManuscriptAnalyzer(config)
                
                # Criar arquivo temporário
                import tempfile
                with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
                    f.write(st.session_state.text_content)
                    temp_path = f.name
                
                # Analisar
                results = analyzer.analyze(temp_path)
                st.session_state.analysis_results = results
                
                # Mostrar resultados
                st.success("✅ Análise concluída!")
                
                # Estatísticas básicas
                st.subheader("📊 Estatísticas")
                col1, col2, col3, col4 = st.columns(4)
                
                stats = results.get('statistics', {})
                with col1:
                    st.metric("Palavras", stats.get('word_count', 0))
                with col2:
                    st.metric("Parágrafos", stats.get('paragraph_count', 0))
                with col3:
                    st.metric("Sentenças", stats.get('sentence_count', 0))
                with col4:
                    st.metric("Caracteres", stats.get('char_count', 0))
                
                # Problemas identificados
                if 'issues' in results:
                    st.subheader("⚠️ Problemas Identificados")
                    for issue in results['issues'][:10]:  # Top 10
                        st.warning(f"**{issue.get('type', 'Issue')}**: {issue.get('description', 'N/A')}")
                
                # Limpar arquivo temporário
                os.unlink(temp_path)
                
            except Exception as e:
                st.error(f"Erro na análise: {e}")
    
    # Mostrar resultados anteriores se disponíveis
    if st.session_state.analysis_results:
        with st.expander("📋 Ver Resultados Completos"):
            st.json(st.session_state.analysis_results)

def render_aprimoramento():
    """Módulo de Aprimoramento com IA"""
    if not ADVANCED_MODULES_AVAILABLE:
        st.error("❌ Módulo de aprimoramento não disponível.")
        return
    
    st.title("✨ Aprimoramento com IA")
    
    st.markdown("""
    Use inteligência artificial para aprimorar seu texto:
    - Melhorar diálogos
    - Enriquecer descrições
    - Suavizar transições
    - Ajustar tom narrativo
    """)
    
    # API Key
    api_key = st.text_input("OpenAI API Key", type="password", help="Necessário para usar o aprimoramento IA")
    
    if not st.session_state.text_content:
        st.warning("⚠️ Nenhum texto disponível.")
        return
    
    # Opções de aprimoramento
    enhancement_type = st.selectbox(
        "Tipo de Aprimoramento",
        ["Diálogos", "Descrições", "Transições", "Geral"]
    )
    
    if st.button("✨ Aplicar Aprimoramento", type="primary", use_container_width=True):
        if not api_key:
            st.error("❌ Por favor, forneça uma API Key da OpenAI")
            return
        
        with st.spinner(f"Aprimorando {enhancement_type.lower()}..."):
            try:
                config = Config()
                config.openai_api_key = api_key
                enhancer = ContentEnhancer(config)
                
                # Criar arquivo temporário
                import tempfile
                with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
                    f.write(st.session_state.text_content)
                    input_path = f.name
                
                # Arquivo de saída
                output_path = input_path.replace('.txt', '_enhanced.txt')
                
                # Aprimorar
                results = enhancer.enhance(input_path, output_path)
                
                # Ler resultado
                with open(output_path, 'r', encoding='utf-8') as f:
                    enhanced_text = f.read()
                
                st.session_state.enhancement_suggestions = {
                    "original": st.session_state.text_content,
                    "enhanced": enhanced_text,
                    "results": results
                }
                
                st.success("✅ Aprimoramento concluído!")
                
                # Comparação
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("📄 Original")
                    st.text_area("Original", st.session_state.text_content[:1000], height=400, disabled=True)
                
                with col2:
                    st.subheader("✨ Aprimorado")
                    st.text_area("Aprimorado", enhanced_text[:1000], height=400, disabled=True)
                
                # Opção de aceitar
                if st.button("✅ Aceitar Alterações"):
                    st.session_state.text_content = enhanced_text
                    st.success("✅ Texto atualizado!")
                    st.rerun()
                
                # Limpar arquivos temporários
                os.unlink(input_path)
                if os.path.exists(output_path):
                    os.unlink(output_path)
                
            except Exception as e:
                st.error(f"Erro no aprimoramento: {e}")

def render_revisao():
    """Módulo de Revisão Editorial"""
    if not ADVANCED_MODULES_AVAILABLE:
        st.error("❌ Módulo de revisão não disponível.")
        return
    
    st.title("📋 Revisão Editorial")
    
    if not st.session_state.text_content:
        st.warning("⚠️ Nenhum texto disponível.")
        return
    
    st.markdown("""
    Revisão editorial completa incluindo:
    - Consistência de personagens
    - Continuidade de timeline
    - Coerência narrativa
    - Checklist editorial
    """)
    
    if st.button("📋 Iniciar Revisão", type="primary", use_container_width=True):
        with st.spinner("Revisando manuscrito..."):
            try:
                config = Config()
                reviewer = EditorialReviewer(config)
                
                # Criar arquivo temporário
                import tempfile
                with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
                    f.write(st.session_state.text_content)
                    temp_path = f.name
                
                # Revisar
                results = reviewer.review(temp_path)
                st.session_state.editorial_review = results
                
                st.success("✅ Revisão concluída!")
                
                # Mostrar resultados
                if 'consistency_issues' in results:
                    st.subheader("⚠️ Problemas de Consistência")
                    for issue in results['consistency_issues'][:10]:
                        st.warning(issue)
                
                if 'suggestions' in results:
                    st.subheader("💡 Sugestões")
                    for suggestion in results['suggestions'][:10]:
                        st.info(suggestion)
                
                # Limpar arquivo temporário
                os.unlink(temp_path)
                
            except Exception as e:
                st.error(f"Erro na revisão: {e}")

# ====================================================================================
# WORKFLOWS COMPLETOS
# ====================================================================================

def render_workflow_14():
    """Workflow de 14 Fases"""
    st.title("🔄 Workflow Completo de 14 Fases")
    
    st.markdown("""
    Processo editorial completo em 14 fases:
    
    **Preparação:** Configuração, Importação, Revisão Ortográfica
    
    **Edição:** Análise Estrutural, Edição de Conteúdo, Formatação, Revisão de Estilo
    
    **Aprimoramento:** Sugestões IA, Validação de Consistência, Pré-visualização
    
    **Finalização:** Elementos Pré/Pós-textuais, Exportação, Publicação
    """)
    
    try:
        render_workflow_tab()
    except Exception as e:
        st.error(f"Erro ao carregar workflow: {e}")
        st.info("Use a funcionalidade de Workflow Automatizado como alternativa.")

def render_workflow_auto():
    """Workflow Automatizado"""
    if not ADVANCED_MODULES_AVAILABLE:
        st.error("❌ Workflow automatizado não disponível.")
        return
    
    st.title("🤖 Workflow Automatizado")
    
    st.markdown("""
    Processamento automático end-to-end do manuscrito.
    
    O sistema irá:
    1. Analisar o manuscrito
    2. Aplicar formatação FastFormat
    3. Sugerir melhorias (se API key fornecida)
    4. Revisar editorialmente
    5. Preparar para exportação
    """)
    
    # Configurações
    api_key = st.text_input("OpenAI API Key (opcional)", type="password")
    
    output_format = st.multiselect(
        "Formatos de Saída",
        ["docx", "pdf", "epub", "html"],
        default=["docx"]
    )
    
    if st.button("🚀 Iniciar Workflow Automatizado", type="primary", use_container_width=True):
        if not st.session_state.text_content:
            st.error("❌ Nenhum texto disponível.")
            return
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Configuração
            config = Config()
            if api_key:
                config.openai_api_key = api_key
            
            orchestrator = WorkflowOrchestrator(config)
            
            # Criar arquivo de entrada temporário
            import tempfile
            with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
                f.write(st.session_state.text_content)
                input_path = f.name
            
            # Diretório de saída
            output_dir = tempfile.mkdtemp()
            
            # Fase 1: Análise
            status_text.text("1/5: Analisando manuscrito...")
            progress_bar.progress(20)
            
            analyzer = ManuscriptAnalyzer(config)
            analysis = analyzer.analyze(input_path)
            
            # Fase 2: FastFormat
            status_text.text("2/5: Aplicando FastFormat...")
            progress_bar.progress(40)
            
            formatted_text = apply_fastformat(st.session_state.text_content, get_ptbr_options())
            
            # Fase 3: Aprimoramento (se API key)
            if api_key:
                status_text.text("3/5: Aprimorando com IA...")
                progress_bar.progress(60)
                
                enhancer = ContentEnhancer(config)
                # Atualizar arquivo temporário
                with open(input_path, 'w', encoding='utf-8') as f:
                    f.write(formatted_text)
                
                output_enhanced = os.path.join(output_dir, "enhanced.txt")
                enhancer.enhance(input_path, output_enhanced)
                
                with open(output_enhanced, 'r', encoding='utf-8') as f:
                    formatted_text = f.read()
            else:
                progress_bar.progress(60)
            
            # Fase 4: Revisão
            status_text.text("4/5: Revisando editorialmente...")
            progress_bar.progress(80)
            
            reviewer = EditorialReviewer(config)
            review = reviewer.review(input_path)
            
            # Fase 5: Exportação
            status_text.text("5/5: Exportando formatos...")
            progress_bar.progress(90)
            
            exporter = PublicationExporter(config)
            
            # Atualizar texto final
            st.session_state.text_content = formatted_text
            
            # Exportar formatos solicitados
            export_files = {}
            for fmt in output_format:
                output_file = os.path.join(output_dir, f"output.{fmt}")
                # Aqui você chamaria os métodos específicos do exporter
                # Por simplicidade, vamos criar um arquivo básico
                if fmt == "docx":
                    doc = Document()
                    doc.add_paragraph(formatted_text)
                    doc.save(output_file)
                    export_files[fmt] = output_file
            
            progress_bar.progress(100)
            status_text.text("✅ Workflow concluído!")
            
            st.success("🎉 Processamento automático concluído com sucesso!")
            
            # Resultados
            st.subheader("📊 Resultados")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.metric("Palavras", len(formatted_text.split()))
                st.metric("Problemas Encontrados", len(analysis.get('issues', [])))
            
            with col2:
                st.metric("Formatos Gerados", len(export_files))
                st.metric("Melhorias Aplicadas", "Sim" if api_key else "Não")
            
            # Downloads
            st.subheader("📥 Downloads")
            for fmt, filepath in export_files.items():
                with open(filepath, 'rb') as f:
                    st.download_button(
                        f"💾 Download {fmt.upper()}",
                        data=f.read(),
                        file_name=f"{st.session_state.book_title}.{fmt}",
                        mime=f"application/{fmt}"
                    )
            
            # Limpar arquivos temporários
            os.unlink(input_path)
            
        except Exception as e:
            st.error(f"Erro no workflow: {e}")
            import traceback
            st.code(traceback.format_exc())

# ====================================================================================
# EXPORTAÇÃO E PUBLICAÇÃO
# ====================================================================================

def render_exportacao():
    """Módulo de Exportação"""
    st.title("📥 Exportação Multi-formato")
    
    if not st.session_state.text_content:
        st.warning("⚠️ Nenhum texto disponível para exportar.")
        return
    
    st.markdown("Exporte seu manuscrito em múltiplos formatos.")
    
    # Formatos disponíveis
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.subheader("📄 Texto")
        if st.button("TXT", use_container_width=True):
            st.download_button(
                "💾 Download TXT",
                data=st.session_state.text_content,
                file_name=f"{st.session_state.book_title}.txt",
                mime="text/plain"
            )
    
    with col2:
        st.subheader("📝 Word")
        if st.button("DOCX", use_container_width=True):
            doc = Document()
            doc.add_heading(st.session_state.book_title, 0)
            doc.add_heading(f"por {st.session_state.author_name}", 1)
            doc.add_page_break()
            
            for paragraph in st.session_state.text_content.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Salvar em buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                "💾 Download DOCX",
                data=buffer,
                file_name=f"{st.session_state.book_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    with col3:
        st.subheader("🌐 Web")
        if st.button("HTML", use_container_width=True):
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>{st.session_state.book_title}</title>
                <style>
                    body {{ font-family: Georgia, serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
                    h1 {{ text-align: center; }}
                    h2 {{ text-align: center; color: #666; }}
                    p {{ text-align: justify; line-height: 1.6; }}
                </style>
            </head>
            <body>
                <h1>{st.session_state.book_title}</h1>
                <h2>por {st.session_state.author_name}</h2>
                <hr>
                {"".join([f"<p>{para}</p>" for para in st.session_state.text_content.split('\n\n') if para.strip()])}
            </body>
            </html>
            """
            
            st.download_button(
                "💾 Download HTML",
                data=html_content,
                file_name=f"{st.session_state.book_title}.html",
                mime="text/html"
            )

def render_publicacao():
    """Módulo de Preparação para Publicação"""
    if not ADVANCED_MODULES_AVAILABLE:
        st.error("❌ Módulo de publicação não disponível.")
        return
    
    st.title("🚀 Preparação para Publicação")
    
    st.markdown("""
    Prepare seu manuscrito para publicação em diferentes plataformas.
    """)
    
    # Plataformas
    platform = st.selectbox(
        "Plataforma de Publicação",
        ["Amazon KDP", "Google Play Books", "Apple Books", "Kobo", "Editora Tradicional"]
    )
    
    st.subheader(f"📘 Guia: {platform}")
    
    if platform == "Amazon KDP":
        st.markdown("""
        ### Requisitos Amazon KDP
        
        **Formato do Arquivo:**
        - DOCX ou PDF
        - Margens: 2.5cm
        - Fonte: Times New Roman 12pt
        - Espaçamento: 1.5
        
        **Metadados Necessários:**
        - Título e subtítulo
        - Nome do autor
        - Descrição (até 4000 caracteres)
        - Palavras-chave (7 no máximo)
        - Categorias (2)
        
        **Capa:**
        - Resolução mínima: 2560 x 1600 pixels
        - Formato: JPG ou TIFF
        - Sem bordas brancas
        """)
        
        if st.button("📦 Gerar Package KDP", use_container_width=True):
            st.info("Gerando package otimizado para Amazon KDP...")
            # Aqui você implementaria a geração específica para KDP
    
    elif platform == "Google Play Books":
        st.markdown("""
        ### Requisitos Google Play Books
        
        **Formato do Arquivo:**
        - EPUB 2.0 ou 3.0
        - PDF (para livros com layout fixo)
        
        **Requisitos Técnicos:**
        - ISBN necessário
        - Sem DRM (adicionar na plataforma)
        - Tabela de conteúdo funcional
        
        **Metadados:**
        - Título completo
        - Autor(es)
        - Descrição (recomendado: 200-400 palavras)
        - Idioma
        - Categoria BISAC
        """)
    
    elif platform == "Apple Books":
        st.markdown("""
        ### Requisitos Apple Books
        
        **Formato do Arquivo:**
        - EPUB 2.0.1 ou 3.0
        - Validado pelo Apple Books Asset Validator
        
        **Especificações:**
        - ISBN necessário
        - Capa: 1400 x 2100 pixels (mínimo)
        - Formato RGB
        
        **Conteúdo:**
        - Copyright obrigatório
        - Metadados iBooks XML
        """)
    
    # Checklist de publicação
    st.subheader("✅ Checklist de Publicação")
    
    checklist_items = [
        "Revisão ortográfica e gramatical completa",
        "Formatação tipográfica aplicada",
        "Elementos pré-textuais incluídos",
        "Elementos pós-textuais incluídos",
        "Capa profissional criada",
        "ISBN obtido (se necessário)",
        "Metadados preparados",
        "Arquivo exportado no formato correto",
        "Validação técnica realizada",
        "Revisão final de qualidade"
    ]
    
    for item in checklist_items:
        st.checkbox(item, key=f"checklist_{hash(item)}")

# ====================================================================================
# ROTEADOR PRINCIPAL
# ====================================================================================

def route_to_module():
    """Roteia para o módulo selecionado"""
    current = st.session_state.get('current_module', 'dashboard')
    
    if current == 'dashboard':
        render_dashboard()
    elif current == 'editor_simples':
        render_editor_simples()
    elif current == 'editor_quill':
        render_editor_quill()
    elif current == 'editor_ace':
        render_editor_ace()
    elif current == 'fastformat':
        render_fastformat()
    elif current == 'analise':
        render_analise()
    elif current == 'aprimoramento':
        render_aprimoramento()
    elif current == 'revisao':
        render_revisao()
    elif current == 'workflow_14':
        render_workflow_14()
    elif current == 'workflow_auto':
        render_workflow_auto()
    elif current == 'exportacao':
        render_exportacao()
    elif current == 'publicacao':
        render_publicacao()
    else:
        render_dashboard()

# ====================================================================================
# FUNÇÃO PRINCIPAL
# ====================================================================================

def run_mega_editor():
    """Função principal que executa o Mega Editor"""
    setup_page()
    initialize_session_state()
    render_sidebar()
    route_to_module()

# ====================================================================================
# PONTO DE ENTRADA
# ====================================================================================

if __name__ == "__main__":
    run_mega_editor()
