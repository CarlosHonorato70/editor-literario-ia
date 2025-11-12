"""
Módulo de Exportação
Exporta o manuscrito em múltiplos formatos e gera documentação.
"""

import os
from pathlib import Path
from typing import Dict, List
from datetime import datetime
import logging

from .config import Config
from .utils import print_info, print_success, write_file_safe

class PublicationExporter:
    """Exporta manuscrito para publicação."""
    
    def __init__(self, config: Config):
        self.config = config
        self.logger = logging.getLogger(__name__)
    
    def export_all(self, formatted_doc: Dict, elements: Dict, metadata: Dict, output_path: Path) -> Dict:
        """
        Exporta todos os arquivos para publicação.
        
        Args:
            formatted_doc: Documento formatado
            elements: Elementos pré e pós-textuais
            metadata: Metadados do manuscrito
            output_path: Diretório de saída
            
        Returns:
            Dicionário com arquivos gerados
        """
        print_info("Preparando pacote de publicação...")
        
        files_generated = []
        
        # 1. Unifica documento completo
        complete_doc = self._unify_document(formatted_doc, elements)
        
        # 2. Exporta em formatos solicitados
        if "md" in self.config.export_formats:
            md_path = output_path / "Livro_Pronto_Para_Publicacao.md"
            if write_file_safe(str(md_path), complete_doc):
                files_generated.append(str(md_path))
                print_success(f"Exportado: {md_path.name}")
        
        if "docx" in self.config.export_formats:
            docx_path = output_path / "Livro_Pronto_Para_Publicacao.docx"
            if self._export_to_docx(complete_doc, str(docx_path)):
                files_generated.append(str(docx_path))
                print_success(f"Exportado: {docx_path.name}")
        
        if "pdf" in self.config.export_formats:
            pdf_path = output_path / "Livro_Pronto_Para_Publicacao.pdf"
            if self._export_to_pdf(complete_doc, str(pdf_path)):
                files_generated.append(str(pdf_path))
                print_success(f"Exportado: {pdf_path.name}")
        
        # 3. Gera metadados de publicação
        if self.config.include_metadata:
            metadata_path = output_path / "Metadados_Publicacao.md"
            metadata_content = self._generate_publication_metadata(metadata, formatted_doc)
            if write_file_safe(str(metadata_path), metadata_content):
                files_generated.append(str(metadata_path))
        
        # 4. Gera guia de próximos passos
        guide_path = output_path / "Guia_Proximos_Passos.md"
        guide_content = self._generate_next_steps_guide(metadata)
        if write_file_safe(str(guide_path), guide_content):
            files_generated.append(str(guide_path))
        
        return {
            "files": files_generated,
            "formats": self.config.export_formats,
            "complete_document_length": len(complete_doc)
        }
    
    def _unify_document(self, formatted_doc: Dict, elements: Dict) -> str:
        """Unifica todos os elementos em documento completo."""
        sections = []
        
        # Elementos pré-textuais
        pre_textual_order = [
            "Folha_Rosto",
            "Ficha_Catalografica",
            "Dedicatoria",
            "Agradecimentos",
            "Prefacio",
            "Sumario"
        ]
        
        for element_name in pre_textual_order:
            if element_name in elements.get("files", {}):
                sections.append(elements["files"][element_name])
                sections.append("\n---\n")
        
        # Conteúdo principal
        sections.append(formatted_doc["content"])
        sections.append("\n---\n")
        
        # Elementos pós-textuais
        post_textual_order = [
            "Referencias",
            "Glossario",
            "Indice_Remissivo"
        ]
        
        for element_name in post_textual_order:
            if element_name in elements.get("files", {}):
                sections.append(elements["files"][element_name])
                sections.append("\n---\n")
        
        return '\n\n'.join(sections)
    
    def _export_to_docx(self, content: str, output_path: str) -> bool:
        """Exporta para formato DOCX."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re
            
            doc = Document()
            
            # Configura estilos
            style = doc.styles['Normal']
            font = style.font
            font.name = self.config.default_font
            font.size = Pt(self.config.default_font_size)
            
            # Processa conteúdo
            lines = content.split('\n')
            
            for line in lines:
                line = line.rstrip()
                
                # Headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                
                # Separadores
                elif line.startswith('---'):
                    doc.add_paragraph('_' * 50)
                
                # Listas
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                
                # Blockquotes
                elif line.startswith('> '):
                    para = doc.add_paragraph(line[2:])
                    para.paragraph_format.left_indent = Inches(0.5)
                
                # Linha vazia
                elif line == '':
                    doc.add_paragraph()
                
                # Parágrafo normal
                else:
                    # Remove marcadores Markdown
                    clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                    clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
                    if clean_line.strip():
                        doc.add_paragraph(clean_line)
            
            doc.save(output_path)
            return True
            
        except ImportError:
            self.logger.warning("python-docx não instalado. Exportação DOCX desabilitada.")
            return False
        except Exception as e:
            self.logger.error(f"Erro ao exportar para DOCX: {e}")
            return False
    
    def _export_to_pdf(self, content: str, output_path: str) -> bool:
        """Exporta para formato PDF."""
        try:
            # Primeiro converte para DOCX temporário
            temp_docx = output_path.replace('.pdf', '_temp.docx')
            if not self._export_to_docx(content, temp_docx):
                return False
            
            # Tenta converter DOCX para PDF
            try:
                from docx2pdf import convert
                convert(temp_docx, output_path)
                os.remove(temp_docx)
                return True
            except ImportError:
                self.logger.warning("docx2pdf não instalado. Usando método alternativo...")
                
                # Método alternativo: usar markdown-pdf
                try:
                    import markdown
                    from weasyprint import HTML
                    
                    # Converte Markdown para HTML
                    html_content = markdown.markdown(content)
                    
                    # Adiciona CSS básico
                    styled_html = f"""
                    <html>
                    <head>
                        <style>
                            body {{
                                font-family: {self.config.default_font};
                                font-size: {self.config.default_font_size}pt;
                                line-height: {self.config.default_line_spacing};
                                margin: 2cm;
                            }}
                            h1 {{ page-break-before: always; }}
                        </style>
                    </head>
                    <body>{html_content}</body>
                    </html>
                    """
                    
                    # Gera PDF
                    HTML(string=styled_html).write_pdf(output_path)
                    os.remove(temp_docx)
                    return True
                    
                except ImportError:
                    self.logger.warning("weasyprint não instalado. Exportação PDF desabilitada.")
                    os.remove(temp_docx)
                    return False
                    
        except Exception as e:
            self.logger.error(f"Erro ao exportar para PDF: {e}")
            return False
    
    def _generate_publication_metadata(self, metadata: Dict, formatted_doc: Dict) -> str:
        """Gera documento de metadados para publicação."""
        title = metadata.get("title", "Título")
        author = metadata.get("author", "Autor")
        word_count = len(formatted_doc["content"].split())
        page_count = word_count // 250
        
        lines = [
            "# METADADOS DE PUBLICAÇÃO",
            "",
            "---",
            "",
            "## INFORMAÇÕES BIBLIOGRÁFICAS",
            "",
            f"**Título:** {title}",
            f"**Autor:** {author}",
            f"**Ano:** {datetime.now().year}",
            "",
            "## ESPECIFICAÇÕES TÉCNICAS",
            "",
            f"**Formato:** {self.config.default_format}",
            f"**Fonte:** {self.config.default_font}",
            f"**Tamanho da Fonte:** {self.config.default_font_size}pt",
            f"**Espaçamento:** {self.config.default_line_spacing}",
            f"**Contagem de Palavras:** {word_count:,}",
            f"**Páginas Estimadas:** {page_count}",
            "",
            "## CLASSIFICAÇÃO",
            "",
            "**Categoria Principal:** [A definir]",
            "**Subcategorias:** [A definir]",
            "**Palavras-chave:** [A definir]",
            "**Público-alvo:** [A definir]",
            "",
            "## SINOPSE",
            "",
            "### Sinopse Curta (até 150 palavras)",
            "",
            "[A ser escrita]",
            "",
            "### Sinopse Longa (até 500 palavras)",
            "",
            "[A ser escrita]",
            "",
            "## INFORMAÇÕES DE PUBLICAÇÃO",
            "",
            "**ISBN:** [A ser obtido]",
            "**Editora:** [A definir]",
            "**Data de Publicação:** [A definir]",
            "**Preço Sugerido:** [A definir]",
            "",
            "## ROTAS DE PUBLICAÇÃO RECOMENDADAS",
            "",
            "### 1. Editora Tradicional",
            "- Maior credibilidade acadêmica",
            "- Distribuição em bibliotecas",
            "- Processo mais longo",
            "",
            "### 2. Autopublicação (Amazon KDP)",
            "- Controle total sobre o processo",
            "- Royalties mais altos",
            "- Lançamento mais rápido",
            "",
            "### 3. Publicação Híbrida",
            "- Combina vantagens de ambas",
            "- Investimento moderado",
            "",
            "---",
            "",
            f"**Documento gerado em:** {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            ""
        ]
        
        return '\n'.join(lines)
    
    def _generate_next_steps_guide(self, metadata: Dict) -> str:
        """Gera guia de próximos passos para publicação."""
        lines = [
            "# GUIA DE PRÓXIMOS PASSOS PARA PUBLICAÇÃO",
            "",
            "---",
            "",
            "## FASE 1: REVISÃO FINAL (1-2 semanas)",
            "",
            "### Ações Necessárias:",
            "",
            "- [ ] Contratar copyeditor profissional",
            "- [ ] Revisar todos os elementos pré-textuais",
            "- [ ] Verificar todas as referências bibliográficas",
            "- [ ] Revisar glossário e índice remissivo",
            "- [ ] Validar formatação em diferentes dispositivos",
            "",
            "### Recursos Necessários:",
            "",
            "- Copyeditor: R$ 3.000 - 8.000",
            "- Tempo: 10-15 dias úteis",
            "",
            "---",
            "",
            "## FASE 2: DESIGN E DIAGRAMAÇÃO (2-3 semanas)",
            "",
            "### Ações Necessárias:",
            "",
            "- [ ] Contratar designer de capa profissional",
            "- [ ] Contratar diagramador para miolo",
            "- [ ] Aprovar design de capa",
            "- [ ] Aprovar diagramação de miolo",
            "- [ ] Gerar arquivos finais (PDF print-ready)",
            "",
            "### Recursos Necessários:",
            "",
            "- Designer de capa: R$ 800 - 3.000",
            "- Diagramação: R$ 1.500 - 5.000",
            "- Tempo: 15-20 dias úteis",
            "",
            "---",
            "",
            "## FASE 3: PREPARAÇÃO PARA PUBLICAÇÃO (1-2 semanas)",
            "",
            "### Ações Necessárias:",
            "",
            "- [ ] Obter ISBN",
            "- [ ] Preparar metadados completos",
            "- [ ] Escrever sinopses (curta e longa)",
            "- [ ] Definir palavras-chave",
            "- [ ] Preparar biografia do autor",
            "- [ ] Criar materiais de divulgação",
            "",
            "### Recursos Necessários:",
            "",
            "- ISBN: R$ 25 - 200 (dependendo da quantidade)",
            "- Copywriter para sinopses: R$ 500 - 1.500 (opcional)",
            "",
            "---",
            "",
            "## FASE 4: PUBLICAÇÃO (1-4 semanas)",
            "",
            "### Rota A: Editora Tradicional",
            "",
            "- [ ] Preparar proposta editorial",
            "- [ ] Enviar para editoras selecionadas",
            "- [ ] Aguardar resposta (2-6 meses)",
            "- [ ] Negociar contrato",
            "- [ ] Processo editorial da editora (6-12 meses)",
            "",
            "### Rota B: Autopublicação (Amazon KDP)",
            "",
            "- [ ] Criar conta KDP",
            "- [ ] Upload de arquivos (capa e miolo)",
            "- [ ] Preencher metadados",
            "- [ ] Definir preço e distribuição",
            "- [ ] Revisar preview",
            "- [ ] Publicar (disponível em 24-72h)",
            "",
            "### Rota C: Publicação Híbrida",
            "",
            "- [ ] Pesquisar editoras híbridas",
            "- [ ] Solicitar orçamentos",
            "- [ ] Escolher pacote de serviços",
            "- [ ] Contratar serviços",
            "- [ ] Acompanhar processo (3-6 meses)",
            "",
            "---",
            "",
            "## FASE 5: LANÇAMENTO E DIVULGAÇÃO (Contínuo)",
            "",
            "### Ações Necessárias:",
            "",
            "- [ ] Criar estratégia de lançamento",
            "- [ ] Preparar materiais de divulgação",
            "- [ ] Contatar influenciadores/resenhistas",
            "- [ ] Organizar evento de lançamento",
            "- [ ] Criar presença online (site, redes sociais)",
            "- [ ] Solicitar resenhas",
            "- [ ] Monitorar vendas e feedback",
            "",
            "### Recursos Necessários:",
            "",
            "- Marketing digital: R$ 1.000 - 5.000/mês",
            "- Evento de lançamento: R$ 2.000 - 10.000",
            "- Website: R$ 500 - 3.000",
            "",
            "---",
            "",
            "## ESTIMATIVAS DE INVESTIMENTO",
            "",
            "### Mínimo (Autopublicação DIY):",
            "- **Total:** R$ 5.000 - 12.000",
            "- **Tempo:** 2-3 meses",
            "",
            "### Recomendado (Autopublicação Profissional):",
            "- **Total:** R$ 12.000 - 25.000",
            "- **Tempo:** 3-4 meses",
            "",
            "### Completo (Publicação Híbrida):",
            "- **Total:** R$ 25.000 - 50.000",
            "- **Tempo:** 6-9 meses",
            "",
            "---",
            "",
            "## RECURSOS ÚTEIS",
            "",
            "### Plataformas de Publicação:",
            "- Amazon KDP (https://kdp.amazon.com)",
            "- Google Play Books",
            "- Apple Books",
            "",
            "### Serviços Profissionais:",
            "- Reedsy (marketplace de profissionais)",
            "- 99designs (design de capa)",
            "- Fiverr/Workana (serviços diversos)",
            "",
            "### Comunidades e Suporte:",
            "- Grupos de autores no Facebook",
            "- Fóruns de autopublicação",
            "- Associações de autores",
            "",
            "---",
            "",
            f"**Guia gerado em:** {datetime.now().strftime('%d/%m/%Y')}",
            "",
            "**Próxima ação recomendada:** Contratar copyeditor profissional para revisão final.",
            ""
        ]
        
        return '\n'.join(lines)


# Standalone function for mega_editor compatibility
def export_document(content, format_type="txt", metadata=None):
    """
    Exporta documento em diferentes formatos.
    Função standalone para compatibilidade com imports diretos.
    
    Args:
        content: Conteúdo do documento
        format_type: Formato de exportação ('txt', 'markdown', 'html')
        metadata: Metadados opcionais do documento
        
    Returns:
        Conteúdo formatado para exportação
    """
    if not content:
        return ""
    
    metadata = metadata or {}
    title = metadata.get('title', 'Documento')
    author = metadata.get('author', '')
    
    if format_type == "txt":
        return content
        
    elif format_type == "markdown":
        header = f"# {title}\n\n"
        if author:
            header += f"**Autor:** {author}\n\n"
        header += "---\n\n"
        return header + content
        
    elif format_type == "html":
        html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
</head>
<body>
    <h1>{title}</h1>
"""
        if author:
            html += f"    <p><strong>Autor:</strong> {author}</p>\n"
        
        html += "    <hr>\n"
        
        # Converte parágrafos
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                html += f"    <p>{para.strip()}</p>\n"
        
        html += """</body>
</html>"""
        return html
    
    return content
