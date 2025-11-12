"""
Módulo de Formatação de Documentos
Realiza formatação e padronização completa do manuscrito.
"""

import re
from typing import Dict, List, Tuple
import logging

from .config import Config
from .utils import print_info, ProgressTracker
from .fastformat_utils import apply_fastformat, get_ptbr_options, get_academic_options
from fastformat import FastFormatOptions

class DocumentFormatter:
    """Formata e padroniza documentos para publicação."""
    
    def __init__(self, config: Config):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self._setup_fastformat_options()
    
    def _setup_fastformat_options(self):
        """Configura opções do FastFormat baseadas na configuração."""
        # Determina o perfil de formatação baseado no tipo de manuscrito
        manuscript_type = getattr(self.config, 'manuscript_type', 'fiction')
        
        if manuscript_type in ['academic', 'technical']:
            self.fastformat_options = get_academic_options()
        else:
            self.fastformat_options = get_ptbr_options()
        
        # Permite configurações customizadas
        self.use_fastformat = getattr(self.config, 'use_fastformat', True)
    
    def format_document(self, enhanced_content: Dict, elements: Dict, corrections: List[Dict]) -> Dict:
        """
        Formata documento completo aplicando todas as padronizações.
        
        Args:
            enhanced_content: Conteúdo aprimorado
            elements: Elementos pré e pós-textuais
            corrections: Correções da revisão editorial
            
        Returns:
            Documento formatado com estatísticas
        """
        print_info("Iniciando formatação e padronização...")
        
        content = enhanced_content["content"]
        
        # Aplica correções da revisão
        content = self._apply_corrections(content, corrections)
        
        # Aplica FastFormat para formatação tipográfica avançada
        if self.use_fastformat:
            print_info("Aplicando FastFormat (formatação tipográfica avançada)...")
            content = apply_fastformat(content, self.fastformat_options)
        
        # Formatação de títulos e subtítulos
        content = self._format_headings(content)
        
        # Formatação de listas
        content = self._format_lists(content)
        
        # Formatação de citações
        content = self._format_quotes(content)
        
        # Formatação de tabelas
        content = self._format_tables(content)
        
        # Formatação de código
        content = self._format_code_blocks(content)
        
        # Padronização de pontuação (se FastFormat não foi usado)
        if not self.use_fastformat:
            content = self._standardize_punctuation(content)
        
        # Padronização de espaçamento
        content = self._standardize_spacing(content)
        
        # Padronização de ênfase
        content = self._standardize_emphasis(content)
        
        # Limpeza final
        content = self._final_cleanup(content)
        
        return {
            "content": content,
            "original_length": len(enhanced_content["content"]),
            "formatted_length": len(content),
            "statistics": {
                "corrections_applied": len(corrections),
                "headings_formatted": content.count('\n#'),
                "lists_formatted": content.count('\n-') + content.count('\n*'),
                "tables_formatted": content.count('|---'),
                "code_blocks_formatted": content.count('```')
            }
        }
    
    def _apply_corrections(self, content: str, corrections: List[Dict]) -> str:
        """Aplica correções identificadas na revisão."""
        for correction in corrections:
            if correction.get("type") == "replacement":
                old_text = correction.get("old")
                new_text = correction.get("new")
                if old_text and new_text:
                    content = content.replace(old_text, new_text)
        
        return content
    
    def _format_headings(self, content: str) -> str:
        """Formata títulos e subtítulos."""
        lines = content.split('\n')
        formatted_lines = []
        
        for line in lines:
            # Garante espaço após # em headings Markdown
            if line.startswith('#'):
                # Remove espaços extras
                line = re.sub(r'^(#+)\s*', r'\1 ', line)
                # Garante que não há # no final
                line = line.rstrip('#').rstrip()
            
            formatted_lines.append(line)
        
        return '\n'.join(formatted_lines)
    
    def _format_lists(self, content: str) -> str:
        """Formata listas."""
        lines = content.split('\n')
        formatted_lines = []
        in_list = False
        
        for i, line in enumerate(lines):
            stripped = line.lstrip()
            
            # Detecta item de lista
            if stripped.startswith(('-', '*', '+')):
                # Padroniza para usar '-'
                indent = len(line) - len(stripped)
                item_text = stripped[1:].lstrip()
                line = ' ' * indent + '- ' + item_text
                in_list = True
            elif stripped.startswith(tuple(f'{n}.' for n in range(10))):
                # Lista numerada - mantém numeração
                in_list = True
            else:
                # Não é item de lista
                if in_list and line.strip() == '':
                    # Linha vazia após lista - mantém
                    in_list = False
            
            formatted_lines.append(line)
        
        return '\n'.join(formatted_lines)
    
    def _format_quotes(self, content: str) -> str:
        """Formata citações (blockquotes)."""
        lines = content.split('\n')
        formatted_lines = []
        
        for line in lines:
            if line.startswith('>'):
                # Garante espaço após >
                line = re.sub(r'^>\s*', '> ', line)
            
            formatted_lines.append(line)
        
        return '\n'.join(formatted_lines)
    
    def _format_tables(self, content: str) -> str:
        """Formata tabelas Markdown."""
        # Padrão para detectar tabelas
        table_pattern = r'(\|.+\|[\r\n]+\|[-:\s|]+\|[\r\n]+(?:\|.+\|[\r\n]+)*)'
        
        def format_table_match(match):
            table = match.group(0)
            lines = table.strip().split('\n')
            
            if len(lines) < 2:
                return table
            
            # Formata cada linha
            formatted_lines = []
            for line in lines:
                # Remove espaços extras ao redor de |
                line = re.sub(r'\s*\|\s*', '|', line)
                # Adiciona espaços ao redor do conteúdo das células
                parts = line.split('|')
                parts = [f' {p.strip()} ' if p.strip() else p for p in parts]
                line = '|'.join(parts)
                formatted_lines.append(line)
            
            return '\n'.join(formatted_lines)
        
        content = re.sub(table_pattern, format_table_match, content, flags=re.MULTILINE)
        
        return content
    
    def _format_code_blocks(self, content: str) -> str:
        """Formata blocos de código."""
        # Garante que blocos de código têm linhas vazias antes e depois
        content = re.sub(r'([^\n])\n```', r'\1\n\n```', content)
        content = re.sub(r'```\n([^\n])', r'```\n\n\1', content)
        
        return content
    
    def _standardize_punctuation(self, content: str) -> str:
        """Padroniza pontuação."""
        # Remove espaços antes de pontuação
        content = re.sub(r'\s+([.,;:!?])', r'\1', content)
        
        # Garante espaço após pontuação (exceto em números decimais)
        content = re.sub(r'([.,;:!?])([A-Za-zÀ-ÿ])', r'\1 \2', content)
        
        # Padroniza reticências
        content = re.sub(r'\.{3,}', '...', content)
        
        # Padroniza travessões
        content = re.sub(r'\s*—\s*', ' — ', content)
        content = re.sub(r'\s*–\s*', ' – ', content)
        
        # Remove espaços múltiplos
        content = re.sub(r' {2,}', ' ', content)
        
        return content
    
    def _standardize_spacing(self, content: str) -> str:
        """Padroniza espaçamento."""
        # Remove linhas vazias excessivas (máximo 2)
        content = re.sub(r'\n{4,}', '\n\n\n', content)
        
        # Garante linha vazia antes de headings
        content = re.sub(r'([^\n])\n(#{1,6} )', r'\1\n\n\2', content)
        
        # Garante linha vazia após headings
        content = re.sub(r'(#{1,6} .+)\n([^\n#])', r'\1\n\n\2', content)
        
        # Remove espaços no final das linhas
        lines = content.split('\n')
        lines = [line.rstrip() for line in lines]
        content = '\n'.join(lines)
        
        return content
    
    def _standardize_emphasis(self, content: str) -> str:
        """Padroniza marcadores de ênfase."""
        # Padroniza negrito para **
        content = re.sub(r'__([^_]+)__', r'**\1**', content)
        
        # Padroniza itálico para *
        content = re.sub(r'_([^_]+)_', r'*\1*', content)
        
        # Remove ênfase vazia
        content = re.sub(r'\*\*\s*\*\*', '', content)
        content = re.sub(r'\*\s*\*', '', content)
        
        return content
    
    def _final_cleanup(self, content: str) -> str:
        """Limpeza final do documento."""
        # Remove linhas vazias no início
        content = content.lstrip('\n')
        
        # Garante que documento termina com uma linha vazia
        content = content.rstrip('\n') + '\n'
        
        # Remove espaços em linhas vazias
        content = re.sub(r'\n\s+\n', '\n\n', content)
        
        return content
    
    def format_markdown_to_docx(self, markdown_content: str, output_path: str) -> bool:
        """
        Converte Markdown formatado para DOCX.
        
        Args:
            markdown_content: Conteúdo em Markdown
            output_path: Caminho do arquivo DOCX de saída
            
        Returns:
            True se sucesso, False se erro
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Configura estilos do documento
            style = doc.styles['Normal']
            font = style.font
            font.name = self.config.default_font
            font.size = Pt(self.config.default_font_size)
            
            # Processa conteúdo linha por linha
            lines = markdown_content.split('\n')
            
            for line in lines:
                line = line.rstrip()
                
                # Headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                
                # Listas
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                
                # Linha vazia
                elif line == '':
                    doc.add_paragraph()
                
                # Parágrafo normal
                else:
                    # Remove marcadores Markdown de ênfase para DOCX
                    clean_line = line
                    clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_line)
                    clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
                    doc.add_paragraph(clean_line)
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            self.logger.error(f"Erro ao converter para DOCX: {e}")
            return False


# Standalone function for mega_editor compatibility
def format_document(content, style="default"):
    """
    Formata o documento com estilo específico.
    Função standalone para compatibilidade com imports diretos.
    
    Args:
        content: Conteúdo do documento
        style: Estilo de formatação ('default', 'compact', 'spaced', etc.)
        
    Returns:
        Conteúdo formatado
    """
    if not content:
        return content
    
    styles = {
        "default": content,
        "compact": content.replace('\n\n', '\n'),
        "spaced": content.replace('\n', '\n\n'),
        "uppercase": content.upper(),
        "lowercase": content.lower(),
        "title": content.title(),
    }
    
    return styles.get(style, content)
