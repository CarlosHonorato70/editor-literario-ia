#!/usr/bin/env python3
"""
Testes para o módulo file_handler.

Este script valida a funcionalidade de extração de texto de arquivos TXT, DOCX e PDF.
"""

import io
import sys
from pathlib import Path
from docx import Document


def print_header(text):
    """Imprime cabeçalho formatado."""
    print("\n" + "=" * 70)
    print(f"  {text}")
    print("=" * 70 + "\n")


def print_success(text):
    """Imprime mensagem de sucesso."""
    print(f"✅ {text}")


def print_error(text):
    """Imprime mensagem de erro."""
    print(f"❌ {text}")


def print_info(text):
    """Imprime mensagem informativa."""
    print(f"ℹ️  {text}")


def test_import_file_handler():
    """Testa importação do módulo file_handler."""
    print_header("TESTE 1: Importação do Módulo file_handler")
    
    try:
        from modules.file_handler import FileHandler, extract_text
        print_success("Módulo file_handler importado com sucesso")
        print_success("Classe FileHandler disponível")
        print_success("Função extract_text disponível")
        return True
    except ImportError as e:
        print_error(f"Falha ao importar file_handler: {e}")
        return False


def test_txt_extraction():
    """Testa extração de texto de arquivo TXT."""
    print_header("TESTE 2: Extração de Texto de TXT")
    
    try:
        from modules.file_handler import extract_text
        
        # Cria conteúdo de teste
        test_content = "Este é um teste de arquivo TXT.\nCom múltiplas linhas.\nE texto em português."
        file_bytes = test_content.encode('utf-8')
        
        # Extrai texto
        text, error = extract_text(file_bytes, "teste.txt")
        
        if error:
            print_error(f"Erro na extração: {error}")
            return False
        
        if text == test_content:
            print_success("Texto TXT extraído corretamente")
            print_info(f"Conteúdo extraído: '{text[:50]}...'")
            return True
        else:
            print_error("Texto extraído não corresponde ao esperado")
            return False
            
    except Exception as e:
        print_error(f"Exceção durante teste TXT: {e}")
        return False


def test_docx_extraction():
    """Testa extração de texto de arquivo DOCX."""
    print_header("TESTE 3: Extração de Texto de DOCX")
    
    try:
        from modules.file_handler import extract_text
        
        # Cria documento DOCX de teste
        doc = Document()
        doc.add_paragraph("Primeiro parágrafo de teste.")
        doc.add_paragraph("Segundo parágrafo de teste.")
        doc.add_paragraph("Terceiro parágrafo em português.")
        
        # Salva em buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        file_bytes = buffer.getvalue()
        
        # Extrai texto
        text, error = extract_text(file_bytes, "teste.docx")
        
        if error:
            print_error(f"Erro na extração: {error}")
            return False
        
        # Verifica se contém os parágrafos
        expected_parts = ["Primeiro parágrafo", "Segundo parágrafo", "Terceiro parágrafo"]
        all_found = all(part in text for part in expected_parts)
        
        if all_found:
            print_success("Texto DOCX extraído corretamente")
            print_info(f"Parágrafos encontrados: {len(expected_parts)}")
            print_info(f"Conteúdo: '{text[:100]}...'")
            return True
        else:
            print_error("Texto extraído não contém todos os parágrafos esperados")
            return False
            
    except Exception as e:
        print_error(f"Exceção durante teste DOCX: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_pdf_extraction():
    """Testa extração de texto de arquivo PDF."""
    print_header("TESTE 4: Extração de Texto de PDF")
    
    try:
        from modules.file_handler import extract_text
        
        # Verifica se PyPDF2 está disponível
        try:
            from PyPDF2 import PdfReader
            print_info("PyPDF2 está disponível")
        except ImportError:
            print_info("PyPDF2 não está instalado - teste será pulado")
            return True  # Não falha se PyPDF2 não estiver instalado
        
        # Cria PDF de teste simples usando reportlab
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            c.drawString(100, 750, "Este é um PDF de teste.")
            c.drawString(100, 730, "Com texto em português.")
            c.save()
            file_bytes = buffer.getvalue()
            
            # Extrai texto
            text, error = extract_text(file_bytes, "teste.pdf")
            
            if error:
                print_error(f"Erro na extração: {error}")
                return False
            
            # Verifica se contém parte do texto
            if "teste" in text.lower():
                print_success("Texto PDF extraído corretamente")
                print_info(f"Conteúdo: '{text[:100]}...'")
                return True
            else:
                print_error("Texto extraído não contém conteúdo esperado")
                print_info(f"Texto extraído: '{text}'")
                return False
                
        except ImportError:
            print_info("reportlab não está disponível - teste será pulado")
            return True  # Não falha se reportlab não estiver disponível
            
    except Exception as e:
        print_error(f"Exceção durante teste PDF: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_unsupported_file():
    """Testa comportamento com arquivo não suportado."""
    print_header("TESTE 5: Arquivo Não Suportado")
    
    try:
        from modules.file_handler import extract_text
        
        # Tenta extrair de arquivo com extensão não suportada
        file_bytes = b"conteudo qualquer"
        text, error = extract_text(file_bytes, "teste.xyz")
        
        if error and "não suportado" in error.lower():
            print_success("Erro apropriado para tipo de arquivo não suportado")
            print_info(f"Mensagem de erro: '{error}'")
            return True
        else:
            print_error("Deveria retornar erro para tipo não suportado")
            return False
            
    except Exception as e:
        print_error(f"Exceção durante teste: {e}")
        return False


def run_all_tests():
    """Executa todos os testes."""
    print_header("TESTES DO MÓDULO FILE_HANDLER")
    
    tests = [
        ("Importação", test_import_file_handler),
        ("Extração TXT", test_txt_extraction),
        ("Extração DOCX", test_docx_extraction),
        ("Extração PDF", test_pdf_extraction),
        ("Arquivo Não Suportado", test_unsupported_file),
    ]
    
    results = []
    for test_name, test_func in tests:
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print_error(f"Erro fatal no teste '{test_name}': {e}")
            results.append((test_name, False))
    
    # Imprime resumo
    print_header("RESUMO DOS TESTES")
    
    passed = sum(1 for _, result in results if result)
    total = len(results)
    
    for test_name, result in results:
        status = "✅ PASSOU" if result else "❌ FALHOU"
        print(f"{status}: {test_name}")
    
    print(f"\n{'='*70}")
    print(f"RESULTADOS: {passed}/{total} testes passaram")
    print(f"{'='*70}\n")
    
    return passed == total


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
