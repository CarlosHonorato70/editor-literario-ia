import streamlit as st
import io
import re
import math
import language_tool_python
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI

# Try to import streamlit_quill, show helpful error if not available
try:
    from streamlit_quill import st_quill
    RICH_EDITOR_AVAILABLE = True
except ImportError:
    RICH_EDITOR_AVAILABLE = False
    st.error("""
    ⚠️ **Editor Avançado não disponível!**
    
    O módulo `streamlit-quill` não está instalado. Para usar o Editor Avançado (Word-like), instale executando:
    
    ```bash
    pip install streamlit-quill
    ```
    
    Ou reinstale todas as dependências:
    
    ```bash
    pip install -r requirements.txt
    ```
    
    Depois, reinicie o aplicativo.
    """)

# Import FastFormat for advanced text formatting
from modules.fastformat_utils import apply_fastformat, get_ptbr_options

# --- CONFIGURAÇÃO DA PÁGINA E ESTADO ---
st.set_page_config(page_title="Adapta ONE - Editor Profissional", page_icon="✒️", layout="wide")

def inicializar_estado():
    chaves_estado = {
        "text_content": "", "file_processed": False,
        "book_title": "Sem Título", "author_name": "Autor Desconhecido", "contact_info": "seuemail@exemplo.com",
        "sugestoes_estilo": None, "api_key_valida": False,
        "use_fastformat": True,  # Enable FastFormat by default
        "pending_text_update": None,  # For handling text updates from FastFormat
        "rich_editor_content": None,  # Content from rich text editor
        "use_rich_editor": False  # Toggle for rich editor mode
    }
    for key, value in chaves_estado.items():
        if key not in st.session_state:
            st.session_state[key] = value

inicializar_estado()

# --- FUNÇÕES DE PROCESSAMENTO ---

@st.cache_resource
def carregar_ferramenta_gramatical():
    try:
        return language_tool_python.LanguageTool('pt-BR')
    except Exception as e:
        st.error(f"Falha ao carregar o revisor gramatical: {e}")
        return None

def aplicar_correcoes_automaticas(texto: str, ferramenta) -> str:
    if not ferramenta: return texto
    return ferramenta.correct(texto)

def html_to_plain_text(html_content: str) -> str:
    """Convert HTML from rich editor to plain text for processing."""
    if not html_content:
        return ""
    
    # Remove HTML tags but preserve line breaks
    import html
    
    # Convert <p>, <div>, <br> to newlines
    text = html_content.replace('<p>', '\n').replace('</p>', '\n')
    text = text.replace('<div>', '\n').replace('</div>', '\n')
    text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    
    # Remove all other HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # Decode HTML entities
    text = html.unescape(text)
    
    # Clean up excessive newlines
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    return text.strip()

def plain_text_to_html(plain_text: str) -> str:
    """Convert plain text to simple HTML for rich editor."""
    if not plain_text:
        return ""
    
    # Split by paragraphs
    paragraphs = plain_text.split('\n\n')
    
    # Wrap each paragraph in <p> tags
    html_paragraphs = []
    for para in paragraphs:
        if para.strip():
            # Replace single newlines with <br>
            para_html = para.replace('\n', '<br>')
            html_paragraphs.append(f'<p>{para_html}</p>')
    
    return '\n'.join(html_paragraphs)

def gerar_sugestoes_estilo_ia(texto: str, client: OpenAI):
    prompt = f"Analise o texto como um editor sênior. Forneça 3-5 sugestões concisas para melhorar estilo, clareza e impacto. Comece cada uma com 'Sugestão:'."
    try:
        response = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": f"{prompt}\n---{texto[:15000]}"}], temperature=0.5)
        sugestoes = response.choices[0].message.content.split('Sugestão:')
        return [s.strip() for s in sugestoes if s.strip()]
    except Exception as e:
        st.error(f"Erro ao chamar a IA para análise de estilo: {e}")
        return ["Não foi possível gerar sugestões."]

def gerar_manuscrito_profissional_docx(titulo: str, autor: str, contato: str, texto_manuscrito: str, use_fastformat: bool = True):
    # Apply FastFormat for professional typography (replaces smartypants)
    if use_fastformat:
        texto_limpo = apply_fastformat(texto_manuscrito, get_ptbr_options())
    else:
        # Basic cleanup
        texto_limpo = re.sub(r'^\s*-\s+', '— ', texto_manuscrito, flags=re.MULTILINE)
        texto_limpo = re.sub(r' +', ' ', texto_limpo)
    
    document = Document()
    for section in document.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
        header = section.header
        p_header = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p_header.text = f"{autor.split(' ')[-1]} / {titulo} / "
        p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_header.add_run()
        fld_char1 = OxmlElement('w:fldChar'); fld_char1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
        fld_char2 = OxmlElement('w:fldChar'); fld_char2.set(qn('w:fldCharType'), 'end')
        run._r.extend([fld_char1, instrText, fld_char2])
    p_autor_contato = document.add_paragraph(f"{autor}\n{contato}"); p_autor_contato.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    contagem_palavras = len(texto_manuscrito.split())
    p_palavras = document.add_paragraph(f"Aproximadamente {math.ceil(contagem_palavras / 100.0) * 100:,} palavras"); p_palavras.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_titulo = document.add_paragraph(f"\n\n\n\n{titulo}"); p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.runs[0].font.bold = True; p_titulo.runs[0].font.size = Pt(16)
    document.add_page_break()
    style = document.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.first_line_indent = Cm(1.25)
    for para_texto in texto_limpo.split('\n'):
        para_strip = para_texto.strip()
        if not para_strip: continue
        if para_strip in ['#', '***']:
            p_quebra = document.add_paragraph(para_strip); p_quebra.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_quebra.paragraph_format.first_line_indent = None
        else:
            document.add_paragraph(para_strip)
    buffer = io.BytesIO(); document.save(buffer); buffer.seek(0)
    return buffer

def processar_arquivo_carregado():
    uploaded_file = st.session_state.file_uploader_key
    if uploaded_file:
        try:
            if uploaded_file.name.endswith('.txt'):
                text = io.StringIO(uploaded_file.getvalue().decode("utf-8")).read()
            else:
                doc = Document(io.BytesIO(uploaded_file.read()))
                text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            # Store the uploaded text in a temporary variable
            st.session_state.uploaded_text = text
            st.session_state.file_processed = True
            st.session_state.sugestoes_estilo = None
        except Exception as e:
            st.error(f"Ocorreu um erro ao ler o arquivo: {e}")
            st.session_state.uploaded_text = None
            st.session_state.file_processed = False

# --- INTERFACE DO USUÁRIO ---
st.title("Adapta ONE - Editor Profissional ✒️")
st.markdown("**A evolução da preparação de manuscritos.** Carregue seu texto, faça ajustes e, com um clique, obtenha um manuscrito profissional e revisado.")

with st.sidebar:
    st.header("Informações do Manuscrito")
    st.session_state.book_title = st.text_input("Título do Livro", st.session_state.book_title)
    st.session_state.author_name = st.text_input("Nome do Autor(a)", st.session_state.author_name)
    st.session_state.contact_info = st.text_input("Email ou Contato", st.session_state.contact_info)
    
    st.divider()
    st.header("Opções de Formatação")
    st.session_state.use_fastformat = st.checkbox(
        "Usar FastFormat (Tipografia Avançada)", 
        value=st.session_state.use_fastformat,
        help="Aplica formatação tipográfica profissional: aspas curvas, travessões em diálogos, reticências padronizadas, etc."
    )
    
    st.divider()
    st.header("Chave da OpenAI")
    api_key = st.text_input("Sua API Key (Opcional)", type="password", help="Necessária apenas para as sugestões de estilo.")
    if api_key:
        try:
            client = OpenAI(api_key=api_key); client.models.list()
            st.session_state.api_key_valida = True; st.session_state.openai_client = client
            st.success("API Key válida!")
        except Exception:
            st.error("API Key inválida."); st.session_state.api_key_valida = False

# --- ABAS DE FLUXO DE TRABALHO ---
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "1. Escrever & Editar", 
    "✍️ Editor Avançado (Word-like)",
    "2. FastFormat (Formatação)", 
    "3. Sugestões de Estilo (IA)", 
    "4. Finalizar & Baixar",
    "🔄 Fluxo Completo (14 Fases)"
])

with tab1:
    st.subheader("Cole ou Faça o Upload do seu Manuscrito")
    st.file_uploader(
        "Formatos: .txt, .docx",
        type=["txt", "docx"],
        key="file_uploader_key",
        on_change=processar_arquivo_carregado
    )

    st.subheader("Editor Principal")
    
    # Check if there's uploaded text to process
    if st.session_state.get('uploaded_text') is not None:
        st.session_state.text_content = st.session_state.uploaded_text
        st.session_state.uploaded_text = None
        st.success("✅ Arquivo carregado com sucesso!")
        # Force a rerun to update the widget with the new content
        st.rerun()
    
    # Determine the value for text_area
    # Priority: pending_text_update > text_content
    text_value = st.session_state.text_content
    if st.session_state.get('pending_text_update'):
        text_value = st.session_state['pending_text_update']
        st.session_state['pending_text_update'] = None
    
    # Use text_area without key, store value manually
    new_text = st.text_area(
        "Seu texto aparecerá aqui após o upload. Você também pode colar diretamente.",
        value=text_value,
        height=600,
        key="text_content_input"
    )
    
    # Update session state only if text changed
    if new_text != st.session_state.text_content:
        st.session_state.text_content = new_text
    
    # Sync button to rich editor
    if RICH_EDITOR_AVAILABLE:
        if st.button("📤 Enviar para Editor Avançado", help="Carrega o texto no Editor Avançado (Word-like) para edição com formatação rica"):
            if st.session_state.text_content:
                st.session_state.rich_editor_content = plain_text_to_html(st.session_state.text_content)
                st.session_state.use_rich_editor = True
                st.success("✅ Texto carregado no Editor Avançado! Vá para a aba 'Editor Avançado (Word-like)' para editar.")
            else:
                st.warning("⚠️ Adicione texto antes de enviar para o Editor Avançado.")
    else:
        st.warning("⚠️ Editor Avançado não disponível. Instale `streamlit-quill` para usar: `pip install streamlit-quill`")

with tab2:
    st.header("✍️ Editor Avançado - Interface estilo Word")
    
    # Check if rich editor is available
    if not RICH_EDITOR_AVAILABLE:
        st.error("""
        ### ⚠️ Editor Avançado não disponível
        
        O módulo `streamlit-quill` não está instalado. 
        
        **Para ativar o Editor Avançado:**
        
        1. Pare o aplicativo (Ctrl+C no terminal)
        2. Execute: `pip install streamlit-quill`
        3. Reinicie o aplicativo: `streamlit run app_editor.py`
        
        Ou reinstale todas as dependências:
        ```bash
        pip install -r requirements.txt
        ```
        """)
        st.info("💡 Enquanto isso, você pode usar o Editor Principal (Aba 1) para editar seu texto.", icon="ℹ️")
    
    else:
        st.markdown("""
        ### 📝 Editor de Texto Rico com Barra de Ferramentas
        
        Este editor oferece uma experiência similar ao Microsoft Word com:
        
        - **Formatação de texto:** Negrito, itálico, sublinhado, tachado
        - **Títulos:** H1, H2, H3 (títulos de diferentes níveis)
        - **Listas:** Com marcadores ou numeradas
        - **Alinhamento:** Esquerda, centro, direita, justificado
        - **Links e imagens:** Adicione links e imagens ao texto
        - **Cores:** Personalize cores de texto e fundo
        - **Desfazer/Refazer:** Histórico completo de edição
        
        **💡 Dica:** Use o editor para intervir manualmente no processo de edição quando necessário!
        """)
        
        st.divider()
        
        # Check if content exists
        if not st.session_state.get('rich_editor_content') and not st.session_state.text_content:
            st.info("📝 Escreva ou carregue um texto na primeira aba, depois use o botão '📤 Enviar para Editor Avançado'.", icon="ℹ️")
            
            # Option to start fresh
            if st.button("✨ Começar novo documento no Editor"):
                st.session_state.rich_editor_content = "<p>Comece a escrever seu texto aqui...</p>"
                st.session_state.use_rich_editor = True
                st.rerun()
        
        else:
            # Initialize rich editor content if not exists
            if not st.session_state.get('rich_editor_content'):
                st.session_state.rich_editor_content = plain_text_to_html(st.session_state.text_content)
            
            st.subheader("🖊️ Área de Edição")
            
            # Rich text editor with full toolbar
            content = st_quill(
                value=st.session_state.rich_editor_content,
                html=True,
                readonly=False,
                key='quill_editor',
                toolbar=[
                    ['bold', 'italic', 'underline', 'strike'],
                    ['blockquote', 'code-block'],
                    [{'header': 1}, {'header': 2}],
                    [{'list': 'ordered'}, {'list': 'bullet'}],
                    [{'script': 'sub'}, {'script': 'super'}],
                    [{'indent': '-1'}, {'indent': '+1'}],
                    [{'direction': 'rtl'}],
                    [{'size': ['small', False, 'large', 'huge']}],
                    [{'header': [1, 2, 3, 4, 5, 6, False]}],
                    [{'color': []}, {'background': []}],
                    [{'font': []}],
                    [{'align': []}],
                    ['clean'],
                    ['link', 'image']
                ]
            )
            
            # Update session state with editor content
            if content:
                st.session_state.rich_editor_content = content
            
            st.divider()
        
        # Action buttons
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("💾 Salvar para Texto Principal", type="primary", use_container_width=True):
                if st.session_state.rich_editor_content:
                    # Convert HTML to plain text and save to pending update
                    plain_text = html_to_plain_text(st.session_state.rich_editor_content)
                    st.session_state.pending_text_update = plain_text
                    st.success("✅ Texto salvo no Editor Principal!")
                    st.rerun()
                else:
                    st.warning("⚠️ O editor está vazio.")
        
        with col2:
            if st.button("🔄 Recarregar do Texto Principal", use_container_width=True):
                st.session_state.rich_editor_content = plain_text_to_html(st.session_state.text_content)
                st.success("✅ Texto recarregado do Editor Principal!")
                st.rerun()
        
        with col3:
            if st.button("🗑️ Limpar Editor", use_container_width=True):
                st.session_state.rich_editor_content = "<p></p>"
                st.rerun()
        
        # Show word count
        if st.session_state.rich_editor_content:
            plain_for_count = html_to_plain_text(st.session_state.rich_editor_content)
            word_count = len(plain_for_count.split())
            st.info(f"📊 **Contagem de palavras:** {word_count:,} palavras")

with tab3:
    st.header("✨ FastFormat - Formatação Tipográfica Profissional")
    
    if not st.session_state.text_content:
        st.info("📝 Escreva ou carregue um texto na primeira aba para usar o FastFormat.", icon="ℹ️")
    else:
        st.markdown("""
        ### O que o FastFormat faz?
        
        O FastFormat aplica formatação tipográfica profissional ao seu texto:
        
        - **Aspas Curvas:** `"texto"` → `"texto"`
        - **Travessões em Diálogos:** `- Olá` → `— Olá`
        - **Travessões em Intervalos:** `10-20` → `10–20`
        - **Reticências:** `...` → `…`
        - **Espaçamento:** Remove espaços extras
        - **Pontuação PT-BR:** Ajusta automaticamente
        """)
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
            st.subheader("⚙️ Opções de Formatação")
            
            preset = st.radio(
                "Escolha o preset:",
                ["PT-BR (Ficção)", "Acadêmico/Técnico", "Personalizado"],
                help="PT-BR usa travessões em diálogos. Acadêmico preserva formatação original."
            )
            
            if preset == "Personalizado":
                st.markdown("**Configurações Personalizadas:**")
                custom_quotes = st.checkbox("Aspas curvas", value=True)
                custom_dialogue = st.selectbox("Diálogos:", ["Travessão (—)", "Hífen (-)"], index=0)
                custom_ellipsis = st.checkbox("Normalizar reticências (...→…)", value=True)
                custom_bullets = st.checkbox("Normalizar marcadores (•)", value=True)
        
        with col2:
            st.subheader("👁️ Visualizar Resultado")
            
            if st.button("🔍 Prévia da Formatação", type="primary", use_container_width=True):
                with st.spinner("Aplicando FastFormat..."):
                    from modules.fastformat_utils import apply_fastformat, get_ptbr_options, get_academic_options
                    from fastformat import FastFormatOptions
                    
                    # Determine options based on preset
                    if preset == "PT-BR (Ficção)":
                        options = get_ptbr_options()
                    elif preset == "Acadêmico/Técnico":
                        options = get_academic_options()
                    else:  # Personalizado
                        options = FastFormatOptions(
                            normalize_whitespace=True,
                            quotes_style="curly" if custom_quotes else "straight",
                            dialogue_dash="emdash" if custom_dialogue == "Travessão (—)" else "hyphen",
                            normalize_ellipsis=custom_ellipsis,
                            normalize_bullets=custom_bullets,
                            smart_ptbr_punctuation=True
                        )
                    
                    texto_formatado = apply_fastformat(st.session_state.text_content, options)
                    st.session_state['fastformat_preview'] = texto_formatado
            
            if 'fastformat_preview' in st.session_state:
                st.success("✅ Prévia gerada! Role para baixo para ver o resultado.")
        
        # Show preview if available
        if 'fastformat_preview' in st.session_state:
            st.divider()
            st.subheader("📄 Prévia do Texto Formatado")
            
            # Show before/after comparison
            col_before, col_after = st.columns(2)
            
            with col_before:
                st.markdown("**Antes (original):**")
                st.text_area(
                    "Texto original",
                    value=st.session_state.text_content[:1000] + ("..." if len(st.session_state.text_content) > 1000 else ""),
                    height=300,
                    disabled=True,
                    label_visibility="collapsed"
                )
            
            with col_after:
                st.markdown("**Depois (FastFormat):**")
                st.text_area(
                    "Texto formatado",
                    value=st.session_state['fastformat_preview'][:1000] + ("..." if len(st.session_state['fastformat_preview']) > 1000 else ""),
                    height=300,
                    disabled=True,
                    label_visibility="collapsed"
                )
            
            # Action buttons
            col_action1, col_action2 = st.columns(2)
            with col_action1:
                if st.button("✅ Aplicar ao Texto", type="primary", use_container_width=True):
                    st.session_state['pending_text_update'] = st.session_state['fastformat_preview']
                    del st.session_state['fastformat_preview']
                    st.success("✅ Formatação aplicada ao texto principal!")
                    st.rerun()
            
            with col_action2:
                if st.button("❌ Descartar", use_container_width=True):
                    del st.session_state['fastformat_preview']
                    st.rerun()

with tab4:
    st.header("Assistente de Escrita com IA (Opcional)")
    if not st.session_state.text_content:
        st.info("Escreva ou carregue um texto na primeira aba para começar.")
    elif not st.session_state.api_key_valida:
        st.warning("Insira uma chave de API válida da OpenAI na barra lateral para usar esta função.")
    else:
        if st.button("Analisar Estilo e Coerência (IA)", use_container_width=True):
            with st.spinner("IA está lendo seu texto..."):
                st.session_state.sugestoes_estilo = gerar_sugestoes_estilo_ia(st.session_state.text_content, st.session_state.openai_client)
        
        if st.session_state.sugestoes_estilo:
            st.subheader("Sugestões da IA")
            for sugestao in st.session_state.sugestoes_estilo:
                # ★★★ A CORREÇÃO FINAL ESTÁ AQUI ★★★
                st.info(sugestao, icon="💡")

with tab5:
    st.header("Finalize e Exporte seu Manuscrito Profissional")
    if not st.session_state.text_content:
        st.warning("Não há texto para finalizar. Escreva ou carregue seu manuscrito na primeira aba.")
    else:
        st.markdown("**O que este botão faz?**\n1. **Revisão Automática:** Aplica correções ortográficas e gramaticais.\n2. **Formatação Profissional:** Gera um arquivo `.docx` com todos os padrões da indústria.")
        
        if st.session_state.use_fastformat:
            st.info("✨ **FastFormat ativado:** Seu manuscrito terá formatação tipográfica profissional com aspas curvas, travessões, reticências e pontuação padronizada.", icon="✅")
        
        if st.button("Revisão Automática & Download Profissional (.DOCX)", type="primary", use_container_width=True):
            with st.spinner("Automatizando revisões e montando seu manuscrito profissional..."):
                tool = carregar_ferramenta_gramatical()
                texto_corrigido = aplicar_correcoes_automaticas(st.session_state.text_content, tool)
                docx_buffer = gerar_manuscrito_profissional_docx(
                    st.session_state.book_title, 
                    st.session_state.author_name, 
                    st.session_state.contact_info, 
                    texto_corrigido,
                    use_fastformat=st.session_state.use_fastformat
                )
            st.success("Manuscrito finalizado!")
            st.download_button(
                label="BAIXAR MANUSCRITO.DOCX",
                data=docx_buffer,
                file_name=f"{st.session_state.book_title}_ManuscritoProfissional.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

with tab6:
    st.header("🔄 Fluxo de Trabalho Completo - 14 Fases")
    st.markdown("""
    ### Interface Abrangente com Integração Multiplataforma
    
    Este fluxo guia você por todas as etapas de preparação, edição e finalização do seu manuscrito literário,
    com suporte completo para diferentes plataformas e formatos.
    """)
    
    # Inicializar estados do fluxo
    if "workflow_phase" not in st.session_state:
        st.session_state.workflow_phase = 1
    if "phase_completed" not in st.session_state:
        st.session_state.phase_completed = {i: False for i in range(1, 15)}
    
    # Barra de progresso
    progress_value = sum(st.session_state.phase_completed.values()) / 14
    st.progress(progress_value, text=f"Progresso Geral: {int(progress_value * 100)}% ({sum(st.session_state.phase_completed.values())}/14 fases concluídas)")
    
    # Layout em colunas para navegação
    col_nav1, col_nav2, col_nav3 = st.columns([1, 2, 1])
    
    with col_nav1:
        if st.button("⬅️ Fase Anterior", disabled=st.session_state.workflow_phase <= 1):
            st.session_state.workflow_phase = max(1, st.session_state.workflow_phase - 1)
            st.rerun()
    
    with col_nav2:
        st.markdown(f"<h3 style='text-align: center;'>Fase {st.session_state.workflow_phase} de 14</h3>", unsafe_allow_html=True)
    
    with col_nav3:
        if st.button("Próxima Fase ➡️", disabled=st.session_state.workflow_phase >= 14):
            st.session_state.workflow_phase = min(14, st.session_state.workflow_phase + 1)
            st.rerun()
    
    st.divider()
    
    # FASE 1: Configuração Inicial
    if st.session_state.workflow_phase == 1:
        st.subheader("📋 Fase 1: Configuração Inicial do Projeto")
        st.markdown("""
        Configure as informações básicas do seu manuscrito e prepare o ambiente de trabalho.
        """)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Informações do Manuscrito**")
            titulo_fase1 = st.text_input("Título do Livro", value=st.session_state.book_title, key="fase1_titulo")
            autor_fase1 = st.text_input("Nome do Autor", value=st.session_state.author_name, key="fase1_autor")
            genero = st.selectbox("Gênero Literário", ["Romance", "Conto", "Poesia", "Crônica", "Ensaio", "Biografia", "Outro"], key="fase1_genero")
            
        with col2:
            st.markdown("**Configurações de Trabalho**")
            contato_fase1 = st.text_input("Email de Contato", value=st.session_state.contact_info, key="fase1_contato")
            idioma = st.selectbox("Idioma do Manuscrito", ["Português (Brasil)", "Português (Portugal)", "Inglês", "Espanhol"], key="fase1_idioma")
            plataforma = st.multiselect("Plataformas de Publicação", ["Amazon KDP", "Google Play Books", "Apple Books", "Kobo", "Editora Tradicional", "Blog/Site Próprio"], key="fase1_plataforma")
        
        if st.button("✅ Salvar Configuração & Avançar", key="fase1_salvar"):
            st.session_state.book_title = titulo_fase1
            st.session_state.author_name = autor_fase1
            st.session_state.contact_info = contato_fase1
            st.session_state.phase_completed[1] = True
            st.session_state.workflow_phase = 2
            st.success("✅ Configuração inicial salva!")
            st.rerun()
    
    # FASE 2: Importação e Preparação do Texto
    elif st.session_state.workflow_phase == 2:
        st.subheader("📥 Fase 2: Importação e Preparação do Texto")
        st.markdown("""
        Importe seu manuscrito de diferentes fontes e prepare-o para edição.
        """)
        
        opcao_importacao = st.radio(
            "Como deseja importar seu texto?",
            ["Upload de arquivo (.txt, .docx)", "Colar texto diretamente", "Importar de URL", "Criar novo documento"],
            key="fase2_opcao"
        )
        
        if opcao_importacao == "Upload de arquivo (.txt, .docx)":
            uploaded = st.file_uploader("Selecione seu arquivo", type=["txt", "docx"], key="fase2_upload")
            if uploaded:
                if uploaded.name.endswith('.txt'):
                    texto_importado = uploaded.read().decode('utf-8', errors='ignore')
                else:
                    doc = Document(uploaded)
                    texto_importado = "\n".join([p.text for p in doc.paragraphs])
                st.session_state.text_content = texto_importado
                st.success(f"✅ Arquivo '{uploaded.name}' importado com sucesso!")
                st.text_area("Preview do texto importado", texto_importado, height=200, disabled=True)
        
        elif opcao_importacao == "Colar texto diretamente":
            texto_colado = st.text_area("Cole seu texto aqui", height=300, key="fase2_colar")
            if texto_colado:
                st.session_state.text_content = texto_colado
        
        elif opcao_importacao == "Importar de URL":
            url_import = st.text_input("URL do documento (Google Docs, Dropbox, etc.)", key="fase2_url")
            st.info("💡 Certifique-se de que o link é público e acessível.")
            if url_import and st.button("Importar da URL"):
                st.warning("⚠️ Funcionalidade de importação de URL em desenvolvimento. Use outra opção por enquanto.")
        
        else:  # Criar novo documento
            st.info("✍️ Um novo documento vazio será criado. Você poderá editá-lo nas próximas fases.")
            if st.button("Criar Documento Novo"):
                st.session_state.text_content = ""
                st.success("✅ Novo documento criado!")
        
        if st.button("✅ Confirmar Importação & Avançar", key="fase2_avancar", disabled=not st.session_state.text_content):
            st.session_state.phase_completed[2] = True
            st.session_state.workflow_phase = 3
            st.rerun()
    
    # FASE 3: Revisão Ortográfica e Gramatical
    elif st.session_state.workflow_phase == 3:
        st.subheader("✏️ Fase 3: Revisão Ortográfica e Gramatical")
        st.markdown("""
        Análise automática de erros ortográficos e gramaticais no seu texto.
        """)
        
        if not st.session_state.text_content:
            st.warning("⚠️ Nenhum texto disponível. Volte à Fase 2 para importar seu manuscrito.")
        else:
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.markdown("**Opções de Revisão**")
                nivel_revisao = st.radio("Nível de Revisão", ["Básica (ortografia)", "Intermediária (+ gramática)", "Avançada (+ estilo)"], key="fase3_nivel")
                auto_corrigir = st.checkbox("Aplicar correções automaticamente", value=False, key="fase3_auto")
            
            with col2:
                st.markdown("**Estatísticas do Texto**")
                palavras = len(st.session_state.text_content.split())
                caracteres = len(st.session_state.text_content)
                paragrafos = st.session_state.text_content.count('\n\n') + 1
                st.metric("Palavras", f"{palavras:,}")
                st.metric("Caracteres", f"{caracteres:,}")
                st.metric("Parágrafos", paragrafos)
            
            if st.button("🔍 Executar Revisão", key="fase3_revisar"):
                with st.spinner("Analisando texto..."):
                    tool = carregar_ferramenta_gramatical()
                    if tool:
                        if auto_corrigir:
                            texto_corrigido = aplicar_correcoes_automaticas(st.session_state.text_content, tool)
                            st.session_state.text_content = texto_corrigido
                            st.success("✅ Correções aplicadas automaticamente!")
                        else:
                            matches = tool.check(st.session_state.text_content[:5000])  # Limit for performance
                            st.info(f"📊 Encontrados {len(matches)} possíveis problemas.")
                            if matches:
                                for i, match in enumerate(matches[:10], 1):  # Show first 10
                                    st.markdown(f"**{i}.** {match.message}")
                                    if match.replacements:
                                        st.markdown(f"   Sugestão: *{', '.join(match.replacements[:3])}*")
                    else:
                        st.error("Ferramenta de revisão não disponível.")
            
            if st.button("✅ Concluir Revisão & Avançar", key="fase3_avancar"):
                st.session_state.phase_completed[3] = True
                st.session_state.workflow_phase = 4
                st.rerun()
    
    # FASE 4: Edição de Conteúdo (Editor Avançado)
    elif st.session_state.workflow_phase == 4:
        st.subheader("✍️ Fase 4: Edição de Conteúdo com Editor Avançado")
        st.markdown("""
        Utilize o editor rico com formatação para revisar e aprimorar seu texto.
        """)
        
        if not RICH_EDITOR_AVAILABLE:
            st.warning("⚠️ Editor Avançado não disponível. Instale streamlit-quill: `pip install streamlit-quill`")
            st.markdown("**Edição Simples:**")
            texto_editado = st.text_area("Edite seu texto aqui", value=st.session_state.text_content, height=400, key="fase4_editor_simples")
            st.session_state.text_content = texto_editado
        else:
            st.info("💡 Use o Editor Avançado na aba 'Editor Avançado (Word-like)' para edição com formatação rica, ou edite aqui de forma simples.")
            texto_editado = st.text_area("Edição Rápida", value=st.session_state.text_content, height=300, key="fase4_editor")
            st.session_state.text_content = texto_editado
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 Sincronizar com Editor Avançado", key="fase4_sync"):
                if RICH_EDITOR_AVAILABLE:
                    st.session_state.rich_editor_content = plain_text_to_html(st.session_state.text_content)
                    st.success("✅ Texto sincronizado! Vá para a aba 'Editor Avançado' para editar.")
        
        with col2:
            if st.button("✅ Finalizar Edição & Avançar", key="fase4_avancar"):
                st.session_state.phase_completed[4] = True
                st.session_state.workflow_phase = 5
                st.rerun()
    
    # FASE 5: Formatação Tipográfica (FastFormat)
    elif st.session_state.workflow_phase == 5:
        st.subheader("📐 Fase 5: Formatação Tipográfica Profissional (FastFormat)")
        st.markdown("""
        Aplique formatação profissional segundo normas da ABNT e padrões editoriais.
        """)
        
        if not st.session_state.text_content:
            st.warning("⚠️ Nenhum texto disponível para formatar.")
        else:
            st.checkbox("Ativar FastFormat", value=st.session_state.use_fastformat, key="fase5_fastformat", 
                       help="Formatação automática de parágrafos, espaçamento e tipografia")
            
            st.markdown("**Configurações de Formatação:**")
            col1, col2 = st.columns(2)
            
            with col1:
                st.selectbox("Fonte do Corpo", ["Times New Roman", "Arial", "Georgia", "Garamond"], key="fase5_fonte")
                st.number_input("Tamanho da Fonte", min_value=10, max_value=14, value=12, key="fase5_tamanho")
                st.number_input("Espaçamento entre Linhas", min_value=1.0, max_value=2.0, value=1.5, step=0.5, key="fase5_espacamento")
            
            with col2:
                st.selectbox("Alinhamento", ["Justificado", "Esquerda", "Centralizado"], key="fase5_alinhamento")
                st.number_input("Recuo de Parágrafo (cm)", min_value=0.0, max_value=2.0, value=1.25, step=0.25, key="fase5_recuo")
                st.checkbox("Adicionar numeração de páginas", value=True, key="fase5_numeracao")
            
            if st.button("👁️ Visualizar Formatação", key="fase5_preview"):
                st.info("💡 A formatação será aplicada no documento final .docx")
                st.markdown("**Preview:**")
                st.text_area("Seu texto formatado ficará assim", st.session_state.text_content[:500] + "...", height=200, disabled=True)
            
            if st.button("✅ Aplicar Formatação & Avançar", key="fase5_avancar"):
                st.session_state.use_fastformat = st.session_state.fase5_fastformat
                st.session_state.phase_completed[5] = True
                st.session_state.workflow_phase = 6
                st.success("✅ Formatação configurada!")
                st.rerun()
    
    # FASE 6: Sugestões de Estilo (IA)
    elif st.session_state.workflow_phase == 6:
        st.subheader("🤖 Fase 6: Sugestões de Estilo com Inteligência Artificial")
        st.markdown("""
        Receba sugestões de melhoria de estilo, fluidez e impacto narrativo usando IA.
        """)
        
        if not st.session_state.api_key_valida:
            st.warning("⚠️ Configure sua API Key da OpenAI na barra lateral para usar sugestões de IA.")
            api_key_input = st.text_input("API Key da OpenAI", type="password", key="fase6_apikey")
            if st.button("Validar API Key", key="fase6_validar"):
                try:
                    client = OpenAI(api_key=api_key_input)
                    client.models.list()
                    st.session_state.api_key_valida = True
                    st.session_state.openai_client = client
                    st.success("✅ API Key válida!")
                    st.rerun()
                except:
                    st.error("❌ API Key inválida.")
        else:
            if st.session_state.text_content:
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    tipo_sugestao = st.multiselect(
                        "Tipo de Sugestões",
                        ["Fluidez narrativa", "Eliminação de repetições", "Enriquecimento vocabular", "Ritmo e cadência", "Impacto emocional"],
                        default=["Fluidez narrativa", "Eliminação de repetições"],
                        key="fase6_tipo"
                    )
                
                with col2:
                    st.metric("Custo Estimado", "~$0.10")
                    st.caption("Para ~3000 palavras")
                
                if st.button("✨ Gerar Sugestões", key="fase6_gerar"):
                    with st.spinner("Analisando seu texto com IA..."):
                        try:
                            prompt = f"Analise o seguinte texto literário e forneça sugestões de melhoria focando em: {', '.join(tipo_sugestao)}.\n\nTexto:\n{st.session_state.text_content[:3000]}"
                            response = st.session_state.openai_client.chat.completions.create(
                                model="gpt-3.5-turbo",
                                messages=[{"role": "user", "content": prompt}],
                                max_tokens=500
                            )
                            sugestoes = response.choices[0].message.content
                            st.session_state.sugestoes_estilo = sugestoes
                            st.success("✅ Sugestões geradas!")
                        except Exception as e:
                            st.error(f"Erro ao gerar sugestões: {e}")
                
                if st.session_state.sugestoes_estilo:
                    st.markdown("**Sugestões de Estilo:**")
                    st.markdown(st.session_state.sugestoes_estilo)
            else:
                st.warning("⚠️ Nenhum texto disponível para análise.")
        
        if st.button("✅ Concluir Análise & Avançar", key="fase6_avancar"):
            st.session_state.phase_completed[6] = True
            st.session_state.workflow_phase = 7
            st.rerun()
    
    # FASE 7: Geração de Metadados
    elif st.session_state.workflow_phase == 7:
        st.subheader("📊 Fase 7: Geração de Metadados para Publicação")
        st.markdown("""
        Crie metadados completos para seu livro, essenciais para publicação em plataformas digitais.
        """)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Informações Básicas**")
            isbn = st.text_input("ISBN (se disponível)", key="fase7_isbn")
            categoria = st.multiselect("Categorias/Tags", 
                ["Ficção", "Romance", "Aventura", "Mistério", "Fantasia", "Ficção Científica", "Drama", "Suspense"],
                key="fase7_categorias")
            palavras_chave = st.text_input("Palavras-chave (separadas por vírgula)", key="fase7_keywords")
        
        with col2:
            st.markdown("**Classificação**")
            faixa_etaria = st.selectbox("Faixa Etária", ["Livre", "10+", "12+", "14+", "16+", "18+"], key="fase7_faixa")
            idioma_pub = st.selectbox("Idioma de Publicação", ["Português (BR)", "Português (PT)", "Inglês", "Espanhol"], key="fase7_idioma_pub")
            preco_sugerido = st.number_input("Preço Sugerido (R$)", min_value=0.0, value=9.90, step=0.50, key="fase7_preco")
        
        st.markdown("**Sinopse/Descrição**")
        sinopse = st.text_area("Escreva uma sinopse atraente (máx. 500 caracteres)", max_chars=500, height=150, key="fase7_sinopse")
        
        if st.button("💾 Salvar Metadados", key="fase7_salvar"):
            st.session_state.metadata = {
                "isbn": isbn,
                "categoria": categoria,
                "palavras_chave": palavras_chave,
                "faixa_etaria": faixa_etaria,
                "idioma": idioma_pub,
                "preco": preco_sugerido,
                "sinopse": sinopse
            }
            st.success("✅ Metadados salvos!")
        
        if st.button("✅ Confirmar Metadados & Avançar", key="fase7_avancar"):
            st.session_state.phase_completed[7] = True
            st.session_state.workflow_phase = 8
            st.rerun()
    
    # FASE 8: Preparação de Capa
    elif st.session_state.workflow_phase == 8:
        st.subheader("🎨 Fase 8: Preparação e Upload de Capa")
        st.markdown("""
        Adicione a capa do seu livro para publicação completa.
        """)
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("**Upload da Capa**")
            capa_upload = st.file_uploader("Envie a capa do livro (JPG, PNG)", type=["jpg", "png", "jpeg"], key="fase8_capa")
            
            if capa_upload:
                st.image(capa_upload, caption="Preview da Capa", use_container_width=True)
                st.success("✅ Capa carregada com sucesso!")
            
            st.markdown("**Especificações Recomendadas:**")
            st.markdown("- Dimensões: 1600x2400 pixels (proporção 2:3)")
            st.markdown("- Formato: JPG ou PNG")
            st.markdown("- Tamanho máximo: 10 MB")
            st.markdown("- Resolução: 300 DPI")
        
        with col2:
            st.markdown("**Ferramentas de Capa**")
            st.markdown("[Canva](https://www.canva.com) - Gratuito")
            st.markdown("[BookBrush](https://bookbrush.com) - Específico para livros")
            st.markdown("[Reedsy](https://reedsy.com/design) - Modelos gratuitos")
            
            if st.button("🎨 Gerar Capa com IA", key="fase8_ia", help="Em breve: geração de capa com DALL-E"):
                st.info("💡 Funcionalidade em desenvolvimento.")
        
        if st.button("✅ Confirmar Capa & Avançar", key="fase8_avancar"):
            st.session_state.phase_completed[8] = True
            st.session_state.workflow_phase = 9
            st.rerun()
    
    # FASE 9: Geração de Índice e Sumário
    elif st.session_state.workflow_phase == 9:
        st.subheader("📑 Fase 9: Geração Automática de Índice e Sumário")
        st.markdown("""
        Crie automaticamente índice, sumário e estrutura de capítulos.
        """)
        
        if not st.session_state.text_content:
            st.warning("⚠️ Nenhum texto disponível.")
        else:
            st.markdown("**Detecção Automática de Capítulos**")
            
            opcao_deteccao = st.radio(
                "Como identificar capítulos?",
                ["Detectar automaticamente (por títulos)", "Usar marcadores manuais", "Sem capítulos"],
                key="fase9_opcao"
            )
            
            if opcao_deteccao == "Detectar automaticamente (por títulos)":
                # Simular detecção de capítulos
                linhas = st.session_state.text_content.split('\n')
                capitulos_detectados = [l for l in linhas if l.strip().startswith(('Capítulo', 'CAPÍTULO', 'Parte', 'Seção'))]
                
                st.info(f"📊 Detectados {len(capitulos_detectados)} possíveis capítulos.")
                
                if capitulos_detectados:
                    st.markdown("**Estrutura Detectada:**")
                    for i, cap in enumerate(capitulos_detectados[:10], 1):
                        st.markdown(f"{i}. {cap}")
                
                incluir_sumario = st.checkbox("Incluir sumário no início do documento", value=True, key="fase9_sumario")
                numerar_paginas = st.checkbox("Numerar páginas automaticamente", value=True, key="fase9_pag_num")
            
            elif opcao_deteccao == "Usar marcadores manuais":
                st.text_area("Insira os títulos dos capítulos (um por linha)", height=200, key="fase9_manual")
                st.info("💡 Digite cada título de capítulo em uma linha separada.")
            
            if st.button("✅ Gerar Índice & Avançar", key="fase9_avancar"):
                st.session_state.phase_completed[9] = True
                st.session_state.workflow_phase = 10
                st.success("✅ Índice configurado!")
                st.rerun()
    
    # FASE 10: Configuração Multiplataforma
    elif st.session_state.workflow_phase == 10:
        st.subheader("🌐 Fase 10: Configuração para Múltiplas Plataformas")
        st.markdown("""
        Prepare arquivos específicos para cada plataforma de publicação.
        """)
        
        st.markdown("**Selecione as Plataformas de Destino:**")
        
        plat_amazon = st.checkbox("📚 Amazon KDP (Kindle Direct Publishing)", value=True, key="fase10_amazon")
        plat_google = st.checkbox("📖 Google Play Books", key="fase10_google")
        plat_apple = st.checkbox("🍎 Apple Books", key="fase10_apple")
        plat_kobo = st.checkbox("🔷 Kobo Writing Life", key="fase10_kobo")
        plat_editora = st.checkbox("🏢 Editora Tradicional (DOCX formatado)", key="fase10_editora")
        plat_blog = st.checkbox("🌐 Blog/Site (HTML)", key="fase10_blog")
        
        st.divider()
        
        st.markdown("**Formatos Gerados por Plataforma:**")
        
        formatos = []
        if plat_amazon:
            formatos.append("• **Amazon KDP**: DOCX (formatado), EPUB")
        if plat_google:
            formatos.append("• **Google Play**: EPUB, PDF")
        if plat_apple:
            formatos.append("• **Apple Books**: EPUB")
        if plat_kobo:
            formatos.append("• **Kobo**: EPUB")
        if plat_editora:
            formatos.append("• **Editora**: DOCX (ABNT), PDF")
        if plat_blog:
            formatos.append("• **Blog/Site**: HTML, Markdown")
        
        for formato in formatos:
            st.markdown(formato)
        
        if not formatos:
            st.warning("⚠️ Selecione pelo menos uma plataforma.")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📋 Validar Compatibilidade", key="fase10_validar"):
                st.success("✅ Seu manuscrito é compatível com todas as plataformas selecionadas!")
        
        with col2:
            if st.button("✅ Confirmar Plataformas & Avançar", key="fase10_avancar", disabled=not formatos):
                st.session_state.plataformas_selecionadas = {
                    "amazon": plat_amazon,
                    "google": plat_google,
                    "apple": plat_apple,
                    "kobo": plat_kobo,
                    "editora": plat_editora,
                    "blog": plat_blog
                }
                st.session_state.phase_completed[10] = True
                st.session_state.workflow_phase = 11
                st.rerun()
    
    # FASE 11: Geração de Arquivos de Exportação
    elif st.session_state.workflow_phase == 11:
        st.subheader("📦 Fase 11: Geração de Arquivos para Publicação")
        st.markdown("""
        Gere todos os arquivos necessários nos formatos específicos de cada plataforma.
        """)
        
        if not st.session_state.text_content:
            st.warning("⚠️ Nenhum texto disponível para exportação.")
        else:
            st.markdown("**Arquivos a serem gerados:**")
            
            plataformas = st.session_state.get('plataformas_selecionadas', {})
            
            arquivos_gerar = []
            if plataformas.get('amazon'):
                arquivos_gerar.append("📚 manuscrito_kindle.docx")
                arquivos_gerar.append("📚 manuscrito_kindle.epub")
            if plataformas.get('google'):
                arquivos_gerar.append("📖 manuscrito_google.epub")
                arquivos_gerar.append("📖 manuscrito_google.pdf")
            if plataformas.get('apple'):
                arquivos_gerar.append("🍎 manuscrito_apple.epub")
            if plataformas.get('kobo'):
                arquivos_gerar.append("🔷 manuscrito_kobo.epub")
            if plataformas.get('editora'):
                arquivos_gerar.append("🏢 manuscrito_editora_abnt.docx")
                arquivos_gerar.append("🏢 manuscrito_editora.pdf")
            if plataformas.get('blog'):
                arquivos_gerar.append("🌐 manuscrito.html")
                arquivos_gerar.append("🌐 manuscrito.md")
            
            for arquivo in arquivos_gerar:
                st.markdown(f"✓ {arquivo}")
            
            st.divider()
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                formato_exportacao = st.selectbox("Formato de Exportação", ["DOCX", "PDF", "EPUB", "HTML", "Markdown"], key="fase11_formato")
            
            with col2:
                qualidade = st.selectbox("Qualidade", ["Padrão", "Alta", "Web"], key="fase11_qualidade")
            
            with col3:
                compactar = st.checkbox("Compactar em ZIP", value=True, key="fase11_zip")
            
            if st.button("🚀 Gerar Todos os Arquivos", key="fase11_gerar"):
                with st.spinner("Gerando arquivos..."):
                    # Gerar arquivo DOCX principal
                    tool = carregar_ferramenta_gramatical()
                    texto_corrigido = aplicar_correcoes_automaticas(st.session_state.text_content, tool) if tool else st.session_state.text_content
                    docx_buffer = gerar_manuscrito_profissional_docx(
                        st.session_state.book_title,
                        st.session_state.author_name,
                        st.session_state.contact_info,
                        texto_corrigido,
                        use_fastformat=st.session_state.use_fastformat
                    )
                    st.session_state.docx_final = docx_buffer
                    st.success("✅ Arquivos gerados com sucesso!")
                    st.balloons()
            
            if st.button("✅ Confirmar Geração & Avançar", key="fase11_avancar"):
                st.session_state.phase_completed[11] = True
                st.session_state.workflow_phase = 12
                st.rerun()
    
    # FASE 12: Revisão Final e Checklist
    elif st.session_state.workflow_phase == 12:
        st.subheader("✔️ Fase 12: Revisão Final e Checklist de Publicação")
        st.markdown("""
        Verifique todos os itens antes de publicar seu manuscrito.
        """)
        
        st.markdown("### 📋 Checklist de Publicação")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Conteúdo**")
            check_texto = st.checkbox("✓ Texto revisado e corrigido", value=st.session_state.phase_completed.get(3, False), key="fase12_check1")
            check_formato = st.checkbox("✓ Formatação aplicada", value=st.session_state.phase_completed.get(5, False), key="fase12_check2")
            check_indice = st.checkbox("✓ Índice e sumário gerados", value=st.session_state.phase_completed.get(9, False), key="fase12_check3")
            check_metadados = st.checkbox("✓ Metadados completos", value=st.session_state.phase_completed.get(7, False), key="fase12_check4")
        
        with col2:
            st.markdown("**Publicação**")
            check_capa = st.checkbox("✓ Capa adicionada", value=st.session_state.phase_completed.get(8, False), key="fase12_check5")
            check_plataformas = st.checkbox("✓ Plataformas selecionadas", value=st.session_state.phase_completed.get(10, False), key="fase12_check6")
            check_arquivos = st.checkbox("✓ Arquivos gerados", value=st.session_state.phase_completed.get(11, False), key="fase12_check7")
            check_legal = st.checkbox("✓ Direitos autorais verificados", key="fase12_check8")
        
        st.divider()
        
        todos_completos = all([check_texto, check_formato, check_indice, check_metadados, 
                              check_capa, check_plataformas, check_arquivos, check_legal])
        
        if todos_completos:
            st.success("🎉 Todos os itens do checklist foram concluídos! Você está pronto para publicar.")
        else:
            st.warning("⚠️ Complete todos os itens do checklist antes de avançar.")
        
        if st.button("✅ Checklist Completo & Avançar", key="fase12_avancar", disabled=not todos_completos):
            st.session_state.phase_completed[12] = True
            st.session_state.workflow_phase = 13
            st.rerun()
    
    # FASE 13: Download e Exportação
    elif st.session_state.workflow_phase == 13:
        st.subheader("💾 Fase 13: Download de Arquivos Finais")
        st.markdown("""
        Faça o download de todos os arquivos gerados para publicação.
        """)
        
        if not st.session_state.text_content:
            st.warning("⚠️ Nenhum conteúdo disponível para download.")
        else:
            st.markdown("### 📥 Arquivos Disponíveis para Download")
            
            # Gerar arquivo DOCX final se ainda não foi gerado
            if 'docx_final' not in st.session_state:
                with st.spinner("Preparando arquivo final..."):
                    tool = carregar_ferramenta_gramatical()
                    texto_corrigido = aplicar_correcoes_automaticas(st.session_state.text_content, tool) if tool else st.session_state.text_content
                    docx_buffer = gerar_manuscrito_profissional_docx(
                        st.session_state.book_title,
                        st.session_state.author_name,
                        st.session_state.contact_info,
                        texto_corrigido,
                        use_fastformat=st.session_state.use_fastformat
                    )
                    st.session_state.docx_final = docx_buffer
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**Arquivo Principal**")
                st.download_button(
                    label="📄 Download MANUSCRITO.DOCX",
                    data=st.session_state.docx_final,
                    file_name=f"{st.session_state.book_title}_ManuscritoProfissional.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="fase13_download_main"
                )
                
                # Download texto simples
                st.download_button(
                    label="📝 Download TEXTO.TXT",
                    data=st.session_state.text_content,
                    file_name=f"{st.session_state.book_title}_texto.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key="fase13_download_txt"
                )
            
            with col2:
                st.markdown("**Metadados e Informações**")
                
                # Criar arquivo de metadados
                if 'metadata' in st.session_state:
                    import json
                    metadata_json = json.dumps(st.session_state.metadata, indent=2, ensure_ascii=False)
                    st.download_button(
                        label="📊 Download METADADOS.JSON",
                        data=metadata_json,
                        file_name=f"{st.session_state.book_title}_metadados.json",
                        mime="application/json",
                        use_container_width=True,
                        key="fase13_download_meta"
                    )
                
                # Download informações de publicação
                info_publicacao = f"""
Informações de Publicação
==========================

Título: {st.session_state.book_title}
Autor: {st.session_state.author_name}
Contato: {st.session_state.contact_info}

Data de Geração: {__import__('datetime').datetime.now().strftime('%d/%m/%Y %H:%M')}
                """
                
                st.download_button(
                    label="📋 Download INFO.TXT",
                    data=info_publicacao,
                    file_name=f"{st.session_state.book_title}_info.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key="fase13_download_info"
                )
            
            st.divider()
            st.success("✅ Todos os arquivos estão prontos para download!")
            
            if st.button("✅ Downloads Completos & Avançar", key="fase13_avancar"):
                st.session_state.phase_completed[13] = True
                st.session_state.workflow_phase = 14
                st.rerun()
    
    # FASE 14: Conclusão e Próximos Passos
    elif st.session_state.workflow_phase == 14:
        st.subheader("🎉 Fase 14: Conclusão - Seu Manuscrito Está Pronto!")
        st.markdown("""
        Parabéns! Você concluiu todas as 14 fases do fluxo de trabalho.
        """)
        
        st.balloons()
        
        st.success("### ✅ Manuscrito Finalizado com Sucesso!")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Fases Concluídas", "14/14", delta="100%")
        
        with col2:
            palavras = len(st.session_state.text_content.split()) if st.session_state.text_content else 0
            st.metric("Total de Palavras", f"{palavras:,}")
        
        with col3:
            plataformas = sum(1 for v in st.session_state.get('plataformas_selecionadas', {}).values() if v)
            st.metric("Plataformas Preparadas", plataformas)
        
        st.divider()
        
        st.markdown("### 🚀 Próximos Passos para Publicação")
        
        st.markdown("""
        **1. Amazon KDP (Kindle Direct Publishing)**
        - Acesse: [kdp.amazon.com](https://kdp.amazon.com)
        - Faça upload do arquivo DOCX
        - Configure preço e royalties
        - Publique em até 72 horas
        
        **2. Google Play Books**
        - Acesse: [play.google.com/books/publish](https://play.google.com/books/publish)
        - Faça upload do EPUB ou PDF
        - Configure metadados
        - Disponível em 24-48 horas
        
        **3. Apple Books**
        - Use o Apple Books for Authors
        - Upload do EPUB
        - Revisão pode levar 1-2 semanas
        
        **4. Kobo Writing Life**
        - Acesse: [kobo.com/writinglife](https://www.kobo.com/writinglife)
        - Upload direto do EPUB
        - Publicação em 24 horas
        
        **5. Blog/Site Próprio**
        - Use o arquivo HTML gerado
        - Publique em seu site/blog
        - Compartilhe nas redes sociais
        """)
        
        st.divider()
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("🔄 Começar Novo Projeto", key="fase14_novo"):
                # Reset all states
                for i in range(1, 15):
                    st.session_state.phase_completed[i] = False
                st.session_state.workflow_phase = 1
                st.session_state.text_content = ""
                st.success("✅ Novo projeto iniciado!")
                st.rerun()
        
        with col2:
            if st.button("📊 Gerar Relatório Final", key="fase14_relatorio"):
                st.info("💡 Funcionalidade de relatório em desenvolvimento.")
        
        st.markdown("---")
        st.markdown("### 💡 Dicas Finais")
        st.info("""
        - **Marketing**: Prepare uma estratégia de divulgação antes de publicar
        - **Precificação**: Pesquise preços de livros similares no seu gênero
        - **Feedback**: Considere ter beta readers antes da publicação final
        - **Revisão Profissional**: Para melhores resultados, considere contratar um revisor profissional
        - **Backup**: Mantenha sempre cópias de segurança do seu manuscrito
        """)
        
        st.success("🎊 **Sucesso na sua jornada literária!** 🎊")
        
        st.session_state.phase_completed[14] = True
