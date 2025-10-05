import os
import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests # Para baixar a imagem da URL
from docx.enum.style import WD_STYLE_TYPE
import time # Para implementar o backoff em caso de Rate Limit

# --- CONFIGURAÇÃO DE CONSTANTES ---

# 1. DICIONÁRIO DE TAMANHOS KDP/GRÁFICA (Miolo)
KDP_SIZES = {
    "Padrão EUA (6x9 in)": {"name": "6 x 9 in", "width_in": 6.0, "height_in": 9.0, "width_cm": 15.24, "height_cm": 22.86, "papel_fator": 0.00115}, # Papel 50lb / 80gsm
    "Padrão A5 (5.83x8.27 in)": {"name": "A5 (14.8 x 21 cm)", "width_in": 5.83, "height_in": 8.27, "width_cm": 14.8, "height_cm": 21.0, "papel_fator": 0.00115},
    "Pocket (5x8 in)": {"name": "5 x 8 in", "width_in": 5.0, "height_in": 8.0, "width_cm": 12.7, "height_cm": 20.32, "papel_fator": 0.00115},
    "Maior (7x10 in)": {"name": "7 x 10 in", "width_in": 7.0, "height_in": 10.0, "width_cm": 17.78, "height_cm": 25.4, "papel_fator": 0.00115},
}

# 2. TEMPLATES DE ESTILO DE DIAGRAMAÇÃO (Ficção e Acadêmico)
STYLE_TEMPLATES = {
    "Romance Clássico (Garamond)": {"font_name": "Garamond", "font_size_pt": 11, "line_spacing": 1.15, "indent": 0.5},
    "Thriller Moderno (Droid Serif)": {"font_name": "Droid Serif", "font_size_pt": 10, "line_spacing": 1.05, "indent": 0.3},
    "Acadêmico/ABNT (Times New Roman 12)": {"font_name": "Times New Roman", "font_size_pt": 12, "line_spacing": 1.5, "indent": 0.0},
}

# 3. CONFIGURAÇÃO DA IA
API_KEY_NAME = "OPENAI_API_KEY"
PROJECT_ID_NAME = "OPENAI_PROJECT_ID"
MODEL_NAME = "gpt-4o-mini" 

# --- INICIALIZAÇÃO DA IA E LAYOUT ---
st.set_page_config(page_title="Editor Pro IA", layout="wide")
st.title("🚀 Editor Pro IA: Publicação Sem Complicações")
st.subheader("Transforme seu manuscrito em um livro profissional, pronto para ABNT e KDP.")

# Variáveis globais para rastrear o status da API
client = None
API_KEY = None
PROJECT_ID = None

try:
    # Tenta carregar as chaves do Streamlit Secrets
    if hasattr(st, 'secrets'):
         if API_KEY_NAME in st.secrets and PROJECT_ID_NAME in st.secrets:
             API_KEY = st.secrets.get(API_KEY_NAME)
             PROJECT_ID = st.secrets.get(PROJECT_ID_NAME)
    
    # Se as chaves estiverem presentes, inicializa o cliente
    if API_KEY and PROJECT_ID:
        client = OpenAI(api_key=API_KEY, project=PROJECT_ID)
    else:
        st.warning(f"Chave e ID do Projeto OpenAI não configurados. A revisão e a geração de capa **NÃO** funcionarão. Por favor, adicione '{API_KEY_NAME}' e '{PROJECT_ID_NAME}' no Streamlit Secrets.")

except Exception as e:
    st.error(f"Erro na inicialização da API: {e}")
    client = None

# Verifica se o cliente foi inicializado com sucesso
is_api_ready = client is not None


# --- FUNÇÕES DE AUXÍLIO ---

def call_openai_api(system_prompt: str, user_content: str, max_tokens: int = 3000, retries: int = 3) -> str:
    """Função genérica para chamar a API da OpenAI com backoff exponencial."""
    
    # Verifica a prontidão da API antes de qualquer chamada
    if not is_api_ready:
        return "[ERRO DE CONEXÃO DA API] Chaves OPENAI_API_KEY e/ou OPENAI_PROJECT_ID não configuradas. Verifique Streamlit Secrets."

    for i in range(retries):
        try:
            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content}
                ],
                temperature=0.7,
                max_tokens=max_tokens
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            if "Rate limit reached" in str(e) and i < retries - 1:
                wait_time = 2 ** i
                st.warning(f"Limite de taxa atingido. Tentando novamente em {wait_time} segundos... (Tentativa {i+1}/{retries})")
                time.sleep(wait_time)
            else:
                st.error(f"Falha ao se comunicar com a OpenAI. Detalhes: {e}")
                return f"[ERRO DE CONEXÃO DA API] Falha: {e}"
    return "[ERRO DE CONEXÃO DA API] Tentativas de conexão esgotadas devido a Rate Limit."


def revisar_paragrafo(paragrafo_texto: str) -> str:
    """Revisão de um único parágrafo."""
    if not paragrafo_texto.strip(): return "" 
    system_prompt = "Você é um editor literário de nível sênior. Sua tarefa é revisar, editar e aprimorar o parágrafo. Corrija gramática, aprimore o estilo e garanta a coerência. Retorne *apenas* o parágrafo revisado, sem comentários."
    user_content = f"Parágrafo a ser editado:\n---\n{paragrafo_texto}\n---"
    texto_revisado = call_openai_api(system_prompt, user_content, max_tokens=500)
    if "[ERRO DE CONEXÃO DA API]" in texto_revisado:
        return paragrafo_texto
    return texto_revisado

def gerar_conteudo_marketing(titulo: str, autor: str, texto_completo: str) -> str:
    """Gera o blurb para contracapa e sugestões de arte."""
    system_prompt = """
    Você é um Copywriter de Best-sellers. Sua tarefa é criar um blurb de contracapa envolvente.
    Gere o resultado *APENAS* com o texto do blurb, sem títulos ou formatação extra.
    """
    user_content = f"Crie um blurb de contracapa de 3-4 parágrafos para este livro: Título: {titulo}, Autor: {autor}. Amostra: {texto_completo[:5000]}"
    return call_openai_api(system_prompt, user_content, max_tokens=1000)

def gerar_relatorio_estrutural(texto_completo: str) -> str:
    """Analisa o texto completo para dar feedback estrutural."""
    system_prompt = "Você é um Editor-Chefe. Gere um breve Relatório de Revisão para o autor. Foque em: Ritmo da Narrativa, Desenvolvimento de Personagens e Estrutura Geral. Use títulos e bullet points."
    user_content = f"MANUSCRITO PARA ANÁLISE:\n---\n{texto_completo[:15000]}\n---"
    return call_openai_api(system_prompt, user_content)

def gerar_elementos_pre_textuais(titulo: str, autor: str, ano: int, texto_completo: str) -> str:
    """Gera o texto de Copyright e a bio para a página Sobre o Autor."""
    system_prompt = """
    Você é um gerente de editora. Gere o conteúdo essencial de abertura e fechamento para um livro (Ficção/Não-Ficção).
    Use os dados fornecidos e o tom do manuscrito para criar uma bio atraente.
    """
    user_content = f"""
    Título: {titulo}
    Autor: {autor}
    Ano: {ano}
    Manuscrito (Amostra): {texto_completo[:5000]}
    
    Gere o resultado no formato estrito:
    
    ### 1. Página de Copyright e Créditos
    
    [Informações sobre Copyright (Ex: Direitos Autorais 2025, Carlos Honorato. Todos os direitos reservados. Proibida a reprodução.)]
    [Informações de Publicação (Ex: Primeira Edição, E-book/Impresso, 2025)]
    [Aviso Legal padrão]
    
    ### 2. Página 'Sobre o Autor'
    
    [Bio envolvente de 2-3 parágrafos, formatada para uma página de livro.]
    """
    return call_openai_api(system_prompt, user_content)

def gerar_relatorio_conformidade_kdp(titulo: str, autor: str, page_count: int, format_data: dict, espessura_cm: float, capa_largura_total_cm: float, capa_altura_total_cm: float) -> str:
    """Gera um checklist de conformidade técnica para upload na Amazon KDP."""
    tamanho_corte = format_data['name']
    prompt_kdp = f"""
    Você é um Especialista Técnico em Publicação e Conformidade da Amazon KDP. Gere um Relatório de Conformidade focado em upload bem-sucedido para Livros Físicos (Brochura) e eBooks.
    
    Gere o relatório usando o formato de lista e títulos, focando em:
    
    ### 1. Livro Físico (Brochura - Especificações)
    - **Tamanho de Corte Final (Miolo):** {tamanho_corte} ({format_data['width_cm']} x {format_data['height_cm']} cm).
    - **Espessura da Lombada (Calculada):** **{espessura_cm} cm**.
    - **Dimensões do Arquivo de Capa (Arte Completa):** **{capa_largura_total_cm} cm (Largura Total) x {capa_altura_total_cm} cm (Altura Total)**.

    ### 2. Checklist de Miolo (DOCX)
    - Verifique se todos os títulos de capítulos estão marcados com o estilo 'Título 1' no DOCX baixado (essencial para Sumário/TOC).
    - As margens internas (lado da lombada) estão em 1.0 polegadas para segurança do corte.
    
    ### 3. Otimização de Metadados (SEO Básico KDP)
    Sugira 3 categorias de nicho da Amazon e 3 palavras-chave de cauda longa para otimizar a listagem do livro.
    """
    return call_openai_api("Você é um especialista em publicação KDP.", prompt_kdp)


# --- FUNÇÕES DOCX AVANÇADAS ---

def adicionar_pagina_rosto(documento: Document, titulo: str, autor: str, style_data: dict):
    """Adiciona uma página de rosto formatada."""
    font_name = style_data['font_name']
    
    documento.add_page_break()

    p_title = documento.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.add_run(titulo).bold = True
    p_title.runs[0].font.size = Pt(24) # Título maior
    p_title.runs[0].font.name = font_name
    
    for _ in range(5):
        documento.add_paragraph()
    
    p_author = documento.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.add_run(autor)
    p_author.runs[0].font.size = Pt(16)
    p_author.runs[0].font.name = font_name

    documento.add_page_break()


def adicionar_pagina_generica(documento: Document, titulo: str, subtitulo: str = None):
    """Adiciona uma página de título formatada e um placeholder."""
    documento.add_page_break()
    
    p_header = documento.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.add_run(titulo).bold = True
    p_header.runs[0].font.size = Pt(18)
    
    if subtitulo:
        p_sub = documento.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.add_run(subtitulo).italic = True
        p_sub.runs[0].font.size = Pt(12)
    
    documento.add_paragraph("")
    
    if titulo == "Sumário":
        p_inst = documento.add_paragraph("⚠️ Para gerar o índice automático, use a função 'Referências' -> 'Sumário' do seu editor de texto. Todos os títulos de capítulo já foram marcados (Estilo: Título 1).")
    else:
        p_inst = documento.add_paragraph("⚠️ Este é um placeholder. Insira o conteúdo real aqui após o download. O espaço e a numeração já estão configurados.")
        
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.runs[0].font.size = Pt(10)
    documento.add_page_break()


# --- FUNÇÃO DE GERAÇÃO DE CAPA COMPLETA ---
def gerar_capa_ia_completa(prompt_visual: str, blurb: str, autor: str, titulo: str, largura_cm: float, altura_cm: float, espessura_cm: float) -> str:
    """
    Chama a API DALL-E 3 para gerar a imagem da capa COMPLETA (Frente, Lombada e Verso).
    """
    
    # Verifica a prontidão da API
    if not is_api_ready:
        return "[ERRO GERAÇÃO DE CAPA] Chaves OPENAI_API_KEY e/ou OPENAI_PROJECT_ID não configuradas. Verifique Streamlit Secrets."
        
    # Prompt avançado para DALL-E 3 para forçar o layout de capa completa
    full_prompt = f"""
    Crie uma imagem de CAPA COMPLETA E ÚNICA para impressão. As dimensões físicas totais (largura x altura) são: {largura_cm} cm x {altura_cm} cm. A lombada tem {espessura_cm} cm de espessura.

    O design deve seguir o estilo: "{prompt_visual}".
    A arte DEVE incluir:
    1. Título '{titulo}' e Autor '{autor}' na capa (Frente).
    2. Título e Autor CLARAMENTE visíveis e centralizados na LOMBADA.
    3. O Blurb de vendas (texto do verso) na CONTRACAPA. Texto: "{blurb[:500]}..." (Use o máximo do texto possível, estilizado).
    4. Crie uma composição coesa que se estenda pela frente, lombada e verso. O design deve ser profissional e pronto para impressão.
    """

    try:
        response = client.images.generate(
            model="dall-e-3",
            prompt=full_prompt,
            size="1792x1024",  # Melhor proporção para capa completa (Horizontal)
            quality="hd", 
            n=1 
        )
        image_url = response.data[0].url
        return image_url
    except Exception as e:
        return f"[ERRO GERAÇÃO DE CAPA] Falha ao gerar a imagem: {e}. Verifique se sua conta OpenAI tem créditos para DALL-E 3 e se o prompt não viola as diretrizes."

# --- FUNÇÃO PRINCIPAL DE DIAGRAMAÇÃO E REVISÃO ---
def processar_manuscrito(uploaded_file, format_data, style_data, incluir_indices_abnt, status_placeholder):
    # is_api_ready é uma variável global
    global is_api_ready 

    documento_original = Document(uploaded_file)
    documento_revisado = Document()
    
    # --- 1. Configuração de Layout (Tamanho de Corte) ---
    section = documento_revisado.sections[0]
    section.page_width = Inches(format_data['width_in'])
    section.page_height = Inches(format_data['height_in'])
    # Margens ajustadas para KDP: 1.0 pol (lombada), 0.6 pol (fora)
    section.left_margin = Inches(1.0) 
    section.right_margin = Inches(0.6) 
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # --- 2. Configuração de Estilo Profissional ---
    font_name = style_data['font_name']
    font_size_pt = style_data['font_size_pt']
    line_spacing = style_data['line_spacing']
    indent_in = style_data['indent']
    
    style = documento_revisado.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size_pt) 
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = line_spacing 
    paragraph_format.first_line_indent = Inches(indent_in) 
    
    # --- 3. Geração dos Elementos Pré-textuais ---
    status_placeholder.info("Fase 1/3: Gerando conteúdo de Copyright e 'Sobre o Autor'...")
    uploaded_file.seek(0)
    manuscript_sample = uploaded_file.getvalue().decode('utf-8', errors='ignore')[:5000]

    if is_api_ready:
        pre_text_content = gerar_elementos_pre_textuais(st.session_state['book_title'], st.session_state['book_author'], 2025, manuscript_sample)
    else:
        pre_text_content = """
        ### 1. Página de Copyright e Créditos
        [Conteúdo de Copyright Padrão: Edite após o download. A IA não está conectada para gerar este texto.]
        
        ### 2. Página 'Sobre o Autor'
        [Conteúdo sobre o Autor Padrão: Edite após o download. A IA não está conectada para gerar a bio envolvente.]
        """

    # --- 4. Inserção de Elementos Pré-textuais no DOCX ---
    
    # Página de Rosto (Título e Autor)
    adicionar_pagina_rosto(documento_revisado, st.session_state['book_title'], st.session_state['book_author'], style_data)
    
    # Página de Copyright
    try:
        # Pega a parte de Copyright, ignorando o cabeçalho '### 2. Página 'Sobre o Autor''
        copyright_text_full = pre_text_content.split('### 2. Página \'Sobre o Autor\'')[0].strip() 
        # Tenta remover o cabeçalho '### 1. Página de Copyright e Créditos' se ele ainda estiver lá
        copyright_text_full = copyright_text_full.replace("### 1. Página de Copyright e Créditos", "").strip() 
    except IndexError:
        copyright_text_full = "[Erro ao extrair o texto de Copyright. Verifique a conexão da API.]"
    
    documento_revisado.add_paragraph("### 1. Página de Copyright e Créditos").bold = True
    p_copy = documento_revisado.add_paragraph(copyright_text_full)
    p_copy.style = 'Normal'
    p_copy.runs[0].font.size = Pt(8) 
    documento_revisado.add_page_break()

    # Sumário (Placeholder e Instruções)
    adicionar_pagina_generica(documento_revisado, "Sumário")
    
    if incluir_indices_abnt:
        adicionar_pagina_generica(documento_revisado, "Índice de Tabelas", "Placeholder para Tabelas")
        adicionar_pagina_generica(documento_revisado, "Índice de Ilustrações", "Placeholder para Ilustrações")

    # --- 5. Processamento do Miolo (Revisão Parágrafo a Parágrafo) ---
    
    paragrafos = documento_original.paragraphs
    total_paragrafos = len(paragrafos)
    texto_completo = ""
    
    progress_bar = st.progress(0, text="Revisando e diagramando o miolo... 0%")
    
    for i, paragrafo in enumerate(paragrafos):
        texto_original = paragrafo.text
        texto_completo += texto_original + "\n"
        
        percent_complete = int((i + 1) / total_paragrafos * 100)
        progress_bar.progress(percent_complete, text=f"Revisando e diagramando o miolo... {percent_complete}%")

        if len(texto_original.strip()) < 10:
            documento_revisado.add_paragraph(texto_original)
            continue 

        # Chama a IA para revisão se a chave estiver configurada
        if is_api_ready:
            texto_revisado = revisar_paragrafo(texto_original)
        else:
            texto_revisado = texto_original # Pula a revisão

        novo_paragrafo = documento_revisado.add_paragraph(texto_revisado)
        
        # Heurística para Títulos de Capítulo (Marcação 'Heading 1' para TOC)
        if len(texto_original.strip()) > 0 and (
            texto_original.strip().lower().startswith("capítulo") or
            texto_original.strip().lower().startswith("introdução") or
            texto_original.strip().lower().startswith("prólogo") or
            texto_original.strip().lower().startswith("conclusão")
        ):
            novo_paragrafo.style = 'Heading 1' 
            novo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            novo_paragrafo.runs[0].font.size = Pt(18) 
            novo_paragrafo.runs[0].font.name = font_name
            documento_revisado.add_paragraph("") 
        else:
            novo_paragrafo.style = 'Normal'
        
    progress_bar.progress(100, text="Fase 2/3: Revisão e diagramação do miolo concluída! 🎉")

    # --- 6. Inserção da Página Pós-Textual ---
    documento_revisado.add_page_break()
    try:
        about_author_text_full = pre_text_content.split('### 2. Página \'Sobre o Autor\'')[1].strip()
        about_author_text_full = about_author_text_full.replace("### 2. Página 'Sobre o Autor'", "").strip() # Remove o cabeçalho
    except IndexError:
        about_author_text_full = "[Erro ao extrair a bio do Autor. Verifique a conexão da API.]"
    
    adicionar_pagina_generica(documento_revisado, "Sobre o Autor", "Sua biografia gerada pela IA")
    # Adiciona a bio do autor
    documento_revisado.add_paragraph(about_author_text_full, style='Normal')

    
    if incluir_indices_abnt:
        adicionar_pagina_generica(documento_revisado, "Apêndice A", "Título do Apêndice")
        adicionar_pagina_generica(documento_revisado, "Anexo I", "Título do Anexo")

    # --- 7. Geração do Blurb de Marketing (para uso na próxima tab) ---
    status_placeholder.info("Fase 3/3: Gerando Blurb de Marketing e preparando para análise...")
    
    if is_api_ready:
        blurb_gerado = gerar_conteudo_marketing(st.session_state['book_title'], st.session_state['book_author'], texto_completo)
    else:
        blurb_gerado = "[Blurb não gerado. Conecte a API para um texto de vendas profissional.]"

    return documento_revisado, texto_completo, blurb_gerado

# --- INICIALIZAÇÃO DE ESTADO ---
if 'book_title' not in st.session_state:
    st.session_state['book_title'] = "O Último Código de Honra"
    st.session_state['book_author'] = "Carlos Honorato"
    st.session_state['page_count'] = 250
    st.session_state['capa_prompt'] = "Um portal antigo se abrindo no meio de uma floresta escura, estilo fantasia épica e mistério, cores roxo e preto, alta resolução."
    st.session_state['blurb'] = "A IA gerará o Blurb (Contracapa) aqui. Edite antes de gerar a capa completa!"
    st.session_state['generated_image_url'] = None
    st.session_state['texto_completo'] = ""
    st.session_state['documento_revisado'] = None
    st.session_state['relatorio_kdp'] = ""
    st.session_state['format_option'] = "Padrão A5 (5.83x8.27 in)"
    st.session_state['incluir_indices_abnt'] = False
    if 'style_option' not in st.session_state:
         st.session_state['style_option'] = "Romance Clássico (Garamond)" # Inicializa com um valor padrão

# --- CÁLCULOS DINÂMICOS (Executados em todas as execuções do script) ---
format_option_default = "Padrão A5 (5.83x8.27 in)"
selected_format_data_calc = KDP_SIZES.get(st.session_state.get('format_option', format_option_default), KDP_SIZES[format_option_default])

espessura_cm = round(st.session_state['page_count'] * selected_format_data_calc['papel_fator'], 2) 
capa_largura_total_cm = round((selected_format_data_calc['width_cm'] * 2) + espessura_cm, 2)
capa_altura_total_cm = round(selected_format_data_calc['height_cm'] + 0.6, 2)
# --- FIM CÁLCULOS DINÂMICOS ---

# --- FLUXO PRINCIPAL DO APLICATIVO (Tabs) ---

config_tab, miolo_tab, capa_tab, export_tab = st.tabs([
    "1. Configuração Inicial", 
    "2. Diagramação & Elementos", 
    "3. Capa Completa IA", 
    "4. Análise & Exportar"
])

# --- TAB 1: CONFIGURAÇÃO INICIAL ---

with config_tab:
    st.subheader("Dados Essenciais para o Projeto")
    
    col1, col2 = st.columns(2)
    with col1:
        st.session_state['book_title'] = st.text_input("Título do Livro", st.session_state['book_title'])
        st.session_state['page_count'] = st.number_input("Contagem Aproximada de Páginas (Miolo)", min_value=10, value=st.session_state['page_count'], step=10, help="Crucial para o cálculo exato da lombada.")
    with col2:
        st.session_state['book_author'] = st.text_input("Nome do Autor", st.session_state['book_author'])
        
    st.subheader("Escolha de Formato e Estilo")
    
    col3, col4, col5 = st.columns(3)
    with col3:
        # Salva o formato escolhido no session_state
        st.session_state['format_option'] = st.selectbox(
            "Tamanho de Corte Final (KDP/Gráfica):",
            options=list(KDP_SIZES.keys()),
            index=list(KDP_SIZES.keys()).index(st.session_state['format_option']), # Mantém o estado
        )
        selected_format_data = KDP_SIZES[st.session_state['format_option']]
    
    with col4:
        # Pega o estilo de volta, ou define o padrão se for a primeira vez
        default_style_key = "Romance Clássico (Garamond)"
        
        # Tenta pegar o valor atual do session state ou usa o default
        current_style_key = st.session_state.get('style_option', default_style_key) 
        
        style_option = st.selectbox(
            "Template de Estilo de Diagramação:",
            options=list(STYLE_TEMPLATES.keys()),
            index=list(STYLE_TEMPLATES.keys()).index(current_style_key), 
            key='style_option', # Usa uma chave de estado para o selectbox
            help="Define fonte, tamanho e espaçamento (Ex: ABNT para trabalhos acadêmicos)."
        )
        selected_style_data = STYLE_TEMPLATES[style_option]
        st.session_state['style_option'] = style_option # Atualiza o estado
        
    with col5:
        incluir_indices_abnt = st.checkbox(
            "Incluir Índices/Apêndices ABNT", 
            value=st.session_state['incluir_indices_abnt'], 
            key='incluir_indices_abnt_checkbox', # Usa uma chave de estado para o checkbox
            help="Adiciona placeholders para Sumário, Índice de Tabelas, Apêndices e Anexos."
        )
        st.session_state['incluir_indices_abnt'] = incluir_indices_abnt # Atualiza o estado
        
    st.subheader("Upload do Manuscrito")
    uploaded_file = st.file_uploader(
        "Carregue o arquivo .docx do seu manuscrito:", 
        type=['docx'],
        help="O processamento de arquivos grandes pode levar alguns minutos. Recomendamos salvar a cada etapa."
    )
    st.session_state['uploaded_file'] = uploaded_file

    st.info(f"**Cálculo da Lombada (Spine):** **{espessura_cm} cm**. **Dimensão Total da Capa:** **{capa_largura_total_cm} cm x {capa_altura_total_cm} cm**. Esses dados serão usados para a Capa IA.")


# --- TAB 2: DIAGRAMAÇÃO & ELEMENTOS ---

with miolo_tab:
    st.header("Fluxo de Diagramação e Revisão com IA")
    
    uploaded_file = st.session_state.get('uploaded_file')

    if uploaded_file is None:
        st.warning("Por favor, carregue um arquivo .docx na aba '1. Configuração Inicial' para começar.")
    else:
        
        if st.button("▶️ Iniciar Processamento do Miolo (Diagramação e Revisão)"):
            if not is_api_ready:
                st.error("Atenção: As chaves OpenAI não estão configuradas. Apenas a diagramação do miolo e a geração do blurb (placeholder) serão realizadas. A revisão da IA será ignorada.")
            
            st.info("Processamento iniciado! Acompanhe o progresso abaixo...")
            status_placeholder = st.empty()
            
            # Recarrega dados de estado
            selected_format_data = KDP_SIZES[st.session_state['format_option']]
            # Garante que selected_style_data é carregado a partir do estado do selectbox
            selected_style_data = STYLE_TEMPLATES[st.session_state.get('style_option', "Romance Clássico (Garamond)")] 
            
            # --- Executa o processamento em passos (Fases 1 a 3) ---
            uploaded_file.seek(0)
            documento_revisado, texto_completo, blurb_gerado = processar_manuscrito(
                uploaded_file, 
                selected_format_data, 
                selected_style_data, 
                st.session_state['incluir_indices_abnt'], 
                status_placeholder
            )
            
            st.session_state['documento_revisado'] = documento_revisado
            st.session_state['texto_completo'] = texto_completo
            st.session_state['blurb'] = blurb_gerado # Atualiza o blurb gerado
            
            status_placeholder.success("Miolo revisado, diagramado e elementos essenciais inseridos! Prossiga para a Capa.")
            st.toast("Miolo Pronto!", icon="✅")
            
        if st.session_state['documento_revisado']:
            selected_style_data = STYLE_TEMPLATES[st.session_state.get('style_option', "Romance Clássico (Garamond)")] 
            st.success(f"Miolo diagramado no formato **{st.session_state['format_option']}** com o estilo **'{selected_style_data['font_name']}'**.")
            
            st.subheader("Intervenção: Blurb da Contracapa")
            st.warning("O Blurb abaixo será usado no design da Capa Completa e no relatório de análise. Edite-o antes de gerar a capa.")
            # Usa a mesma chave para manter o estado
            st.session_state['blurb'] = st.text_area("Texto de Vendas (Blurb):", st.session_state['blurb'], height=300, key='blurb_text_area')


# --- TAB 3: CAPA COMPLETA IA ---

with capa_tab:
    st.header("Criação da Capa Completa (Frente, Lombada e Verso)")
    
    if st.session_state['texto_completo'] == "":
         st.warning("Por favor, execute o processamento do Miolo (Aba 2) para garantir que o Blurb e o Título estejam prontos.")
    else:
        
        st.subheader("Passo 1: Defina o Conteúdo Visual e de Texto")
        
        st.info(f"O Blurb atual (Contracapa) é: **{st.session_state['blurb'][:150]}...**")
        st.text_input("Título para Capa", st.session_state['book_title'], disabled=True)
        st.text_input("Autor para Capa", st.session_state['book_author'], disabled=True)
        
        st.session_state['capa_prompt'] = st.text_area(
            "Descrição VISUAL da Capa (Estilo DALL-E 3):", 
            st.session_state['capa_prompt'], 
            height=200,
            key='capa_prompt_area',
            help="Descreva a arte que deve aparecer, o estilo (ex: óleo, arte digital, surrealismo) e as cores predominantes."
        )

        st.subheader("Passo 2: Geração")
        st.warning(f"Atenção: O modelo DALL-E 3 gera uma imagem do design COMPLETO com as dimensões de **{capa_largura_total_cm}cm x {capa_altura_total_cm}cm** (calculado).")

        if st.button("🎨 Gerar Capa COMPLETA com IA"):
            if not is_api_ready:
                st.error("Chaves OpenAI não configuradas. Não é possível gerar a imagem.")
            else:
                with st.spinner("Gerando design completo da capa (Frente, Lombada e Verso)... Este processo usa DALL-E 3 e pode levar até um minuto."):
                    image_output = gerar_capa_ia_completa(
                        st.session_state['capa_prompt'],
                        st.session_state['blurb'],
                        st.session_state['book_author'],
                        st.session_state['book_title'],
                        capa_largura_total_cm,
                        capa_altura_total_cm,
                        espessura_cm
                    )
                    
                    if "http" in image_output: 
                        st.session_state['generated_image_url'] = image_output
                        st.success("Design de Capa Completa gerado! Prossiga para a aba '4. Exportar'.")
                    else:
                        st.error(image_output)

        if st.session_state['generated_image_url']:
            st.subheader("Pré-visualização da Capa Gerada")
            st.image(st.session_state['generated_image_url'], caption="Capa Completa (Frente, Lombada e Verso)", use_column_width=True)
            st.info("Lembre-se: Esta é uma imagem do design. As dimensões exatas de impressão estão na aba 'Exportar'.")


# --- TAB 4: ANÁLISE & EXPORTAR ---

with export_tab:
    st.header("Relatórios Finais e Exportação")

    if not st.session_state.get('documento_revisado'):
        st.warning("Por favor, execute o processamento do Miolo (Aba 2) antes de exportar.")
    else:
        
        # is_api_ready é uma variável global

        # --- Relatório Estrutural ---
        st.subheader("1. Relatório Estrutural (Editor-Chefe)")
        if is_api_ready:
            # Geração do Relatório Estrutural (se ainda não gerado)
            if 'relatorio_estrutural' not in st.session_state or st.button("Gerar/Atualizar Relatório Estrutural"):
                 with st.spinner("Analisando ritmo e personagens..."):
                    relatorio = gerar_relatorio_estrutural(st.session_state['texto_completo'])
                    st.session_state['relatorio_estrutural'] = relatorio
            
            if st.session_state.get('relatorio_estrutural') and "[ERRO DE CONEXÃO DA API]" not in st.session_state['relatorio_estrutural']:
                 st.markdown(st.session_state['relatorio_estrutural'])
            elif st.session_state.get('relatorio_estrutural'):
                 st.error(st.session_state['relatorio_estrutural']) # Mostra o erro da API se não conseguir gerar
            else:
                 st.info("Clique no botão acima para gerar o Relatório Estrutural.")

        else:
            st.warning("Relatório Estrutural não gerado. Conecte a API para receber o feedback do Editor-Chefe.")
        
        # --- Relatório KDP/Técnico ---
        if st.button("Gerar/Atualizar Relatório Técnico KDP"):
            if is_api_ready:
                with st.spinner("Gerando checklist técnico e de SEO para o upload..."):
                    relatorio_kdp = gerar_relatorio_conformidade_kdp(
                        st.session_state['book_title'], st.session_state['book_author'], st.session_state['page_count'], selected_format_data_calc, 
                        espessura_cm, capa_largura_total_cm, capa_altura_total_cm
                    )
                    st.session_state['relatorio_kdp'] = relatorio_kdp
                    st.success("Relatório KDP atualizado.")
            else:
                st.error("Conecte a API para gerar o Relatório KDP.")
        
        if st.session_state['relatorio_kdp'] and "[ERRO DE CONEXÃO DA API]" not in st.session_state['relatorio_kdp']:
            st.subheader("2. Relatório de Conformidade KDP (Amazon)")
            st.markdown(st.session_state['relatorio_kdp'])
        elif st.session_state.get('relatorio_kdp'):
             st.error(st.session_state['relatorio_kdp']) # Mostra o erro da API se não conseguir gerar

        # --- Downloads Finais ---
        st.subheader("3. Exportar Produtos Finais")
        
        st.markdown(f"""
        #### 📄 Especificações Técnicas (Valores FINAIS)
        - **Miolo (PDF) Tamanho:** {selected_format_data_calc['width_cm']} cm x {selected_format_data_calc['height_cm']} cm
        - **Espessura da Lombada (Spine):** **{espessura_cm} cm**
        - **Capa Completa (Largura x Altura):** **{capa_largura_total_cm} cm x {capa_altura_total_cm} cm** (Entregue essa dimensão ao designer ou use-a como guia para a arte gerada!)
        """)

        # Download do DOCX Diagramado
        buffer = BytesIO()
        st.session_state['documento_revisado'].save(buffer)
        buffer.seek(0)
        st.download_button(
            label="💾 Baixar Manuscrito Diagramado e Formatado (.docx)",
            data=buffer,
            file_name=f"{st.session_state['book_title'].replace(' ', '_')}_MIOLO_PRO.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Download da Capa Completa
        if st.session_state['generated_image_url']:
            try:
                response = requests.get(st.session_state['generated_image_url'])
                if response.status_code == 200:
                    st.download_button(
                        label="💾 Baixar Design da Capa COMPLETA (.png)",
                        data=response.content,
                        file_name=f"Capa_COMPLETA_{st.session_state['book_title'].replace(' ', '_')}.png",
                        mime="image/png"
                    )
            except Exception:
                st.warning("Não foi possível carregar a imagem da capa para download. Tente gerar novamente.")
        
        st.balloons()
