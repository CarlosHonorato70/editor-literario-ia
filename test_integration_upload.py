#!/usr/bin/env python3
"""
Teste de Integração para o fluxo completo de upload de arquivo.

Este script simula o fluxo de upload de arquivo e verifica se o texto
é corretamente extraído e armazenado no session state.
"""

import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter


def print_header(text):
    """Imprime cabeçalho formatado."""
    print("\n" + "=" * 70)
    print(f"  {text}")
    print("=" * 70 + "\n")


def print_success(text):
    """Imprime mensagem de sucesso."""
    print(f"✅ {text}")


def print_error(text):
    """Imprime mensagem de erro."""
    print(f"❌ {text}")


def print_info(text):
    """Imprime mensagem informativa."""
    print(f"ℹ️  {text}")


def test_integration_flow():
    """Testa o fluxo completo de upload e extração."""
    print_header("TESTE DE INTEGRAÇÃO: Fluxo Completo de Upload")
    
    from modules.file_handler import extract_text
    
    # Teste 1: TXT
    print_info("Teste 1: Arquivo TXT")
    txt_content = "Este é um texto de teste.\nCom múltiplas linhas."
    txt_bytes = txt_content.encode('utf-8')
    text, error = extract_text(txt_bytes, "documento.txt")
    
    if error:
        print_error(f"Falha TXT: {error}")
        return False
    
    if text == txt_content:
        print_success("TXT: Texto extraído corretamente e pronto para session_state")
    else:
        print_error("TXT: Texto não corresponde")
        return False
    
    # Teste 2: DOCX
    print_info("\nTeste 2: Arquivo DOCX")
    doc = Document()
    doc.add_paragraph("Parágrafo 1")
    doc.add_paragraph("Parágrafo 2")
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    
    text, error = extract_text(docx_bytes, "documento.docx")
    
    if error:
        print_error(f"Falha DOCX: {error}")
        return False
    
    if "Parágrafo 1" in text and "Parágrafo 2" in text:
        print_success("DOCX: Texto extraído corretamente e pronto para session_state")
    else:
        print_error("DOCX: Texto não contém parágrafos esperados")
        return False
    
    # Teste 3: PDF
    print_info("\nTeste 3: Arquivo PDF")
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.drawString(100, 750, "Texto PDF de teste")
    c.save()
    pdf_bytes = buffer.getvalue()
    
    text, error = extract_text(pdf_bytes, "documento.pdf")
    
    if error:
        print_error(f"Falha PDF: {error}")
        return False
    
    if "teste" in text.lower():
        print_success("PDF: Texto extraído corretamente e pronto para session_state")
    else:
        print_error(f"PDF: Texto não contém conteúdo esperado. Texto: '{text}'")
        return False
    
    # Teste 4: Fluxo de erro
    print_info("\nTeste 4: Tratamento de Erro")
    text, error = extract_text(b"conteudo", "arquivo.xyz")
    
    if error and not text:
        print_success("Erro: Tratamento de erro funciona corretamente")
    else:
        print_error("Erro: Deveria retornar erro para tipo não suportado")
        return False
    
    return True


def test_session_state_integration():
    """Verifica que o código está preparado para integração com Streamlit."""
    print_header("TESTE: Integração com Session State")
    
    print_info("Verificando estrutura do código...")
    
    # Lê o arquivo app_editor.py
    with open('app_editor.py', 'r', encoding='utf-8') as f:
        app_code = f.read()
    
    # Verifica imports necessários
    checks = [
        ('from modules.file_handler import extract_text', 'Import do file_handler'),
        ('st.session_state.text_content = text', 'Atualização do session_state'),
        ('type=["txt", "docx", "pdf"]', 'Suporte a PDF no file_uploader'),
        ('key="text_content"', 'Vinculação do text_area ao session_state'),
    ]
    
    all_passed = True
    for check_str, description in checks:
        if check_str in app_code:
            print_success(f"{description}: OK")
        else:
            print_error(f"{description}: NÃO ENCONTRADO")
            all_passed = False
    
    return all_passed


def main():
    """Executa todos os testes de integração."""
    print_header("SUITE DE TESTES DE INTEGRAÇÃO")
    
    results = []
    
    # Teste 1: Fluxo completo
    try:
        result = test_integration_flow()
        results.append(("Fluxo Completo de Upload", result))
    except Exception as e:
        print_error(f"Erro fatal no teste de fluxo: {e}")
        import traceback
        traceback.print_exc()
        results.append(("Fluxo Completo de Upload", False))
    
    # Teste 2: Integração com session_state
    try:
        result = test_session_state_integration()
        results.append(("Integração Session State", result))
    except Exception as e:
        print_error(f"Erro fatal no teste de integração: {e}")
        import traceback
        traceback.print_exc()
        results.append(("Integração Session State", False))
    
    # Resumo
    print_header("RESUMO DOS TESTES DE INTEGRAÇÃO")
    
    passed = sum(1 for _, result in results if result)
    total = len(results)
    
    for test_name, result in results:
        status = "✅ PASSOU" if result else "❌ FALHOU"
        print(f"{status}: {test_name}")
    
    print(f"\n{'='*70}")
    print(f"RESULTADOS: {passed}/{total} testes passaram")
    print(f"{'='*70}\n")
    
    if passed == total:
        print_success("TODOS OS TESTES PASSARAM!")
        print_info("O sistema está pronto para processar uploads de arquivos TXT, DOCX e PDF")
        return True
    else:
        print_error("ALGUNS TESTES FALHARAM")
        return False


if __name__ == "__main__":
    import sys
    success = main()
    sys.exit(0 if success else 1)
